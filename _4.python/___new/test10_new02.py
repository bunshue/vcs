"""


"""

import os
import sys
import time
import random
import numpy as np
import pandas as pd

print("------------------------------------------------------------")  # 60個


# chap8-01
# Cars基礎類別
class Cars:
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        self.__factory = factory  # 廠牌屬性(保護)
        self.__model = model  # 型號屬性(保護)
        self.color = color  # 顏色屬性(可改)
        self.seat = seat  # 座位屬性(可改)
        self.__displacement = displacement  # 排氣量或電功率屬性(保護)

    # 屬性(Attribute)
    # 函式唯讀
    # 傳回製造商出廠資料
    def manufacture(self):
        if self.__model.find("電能") >= 0:
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "HP"
            )
        else:
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "CC"
            )

    # 傳回排氣量或電功率(保護)
    def displacement(self):
        return self.__displacement

    # 傳回燃料稅(因為汽油、柴油、電能車計算方式不同，交由各子類別自行定義函式)
    def fueltax(self):
        pass

    # 傳回牌照稅(因為主要的汽油、柴油都是以排氣量計算，只有電能車不同，另外覆寫即可)
    def licensetax(self):
        if self.__displacement <= 500:
            return 1620
        elif self.__displacement <= 600:
            return 2160
        elif self.__displacement <= 1200:
            return 4320
        elif self.__displacement <= 1800:
            return 7120
        elif self.__displacement <= 2400:
            return 11230
        elif self.__displacement <= 3000:
            return 15210
        elif self.__displacement <= 3600:
            return 28220
        elif self.__displacement <= 4200:
            return 28220
        elif self.__displacement <= 4800:
            return 46170
        else:
            return 46170

    # 傳回全部稅金
    def totaltax(self):
        return self.fueltax() + self.licensetax()

    # 方法(Method)
    # 輸出汽車屬性
    def info(self):
        if self.__model.find("電能") >= 0:
            print(
                self.__factory
                + "-"
                + self.__model
                + "-"
                + self.color
                + "色-"
                + str(self.seat)
                + "人座"
                + str(self.__displacement)
                + "HP"
            )
        else:
            print(
                self.__factory
                + "-"
                + self.__model
                + "-"
                + self.color
                + "色-"
                + str(self.seat)
                + "人座"
                + str(self.__displacement)
                + "CC"
            )


# 汽油車衍生類別
class GasolineCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "汽油" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回燃料稅
    def fueltax(self):
        if self.displacement() <= 500:
            return 2160
        elif self.displacement() <= 600:
            return 2880
        elif self.displacement() <= 1200:
            return 4320
        elif self.displacement() <= 1800:
            return 4800
        elif self.displacement() <= 2400:
            return 6180
        elif self.displacement() <= 3000:
            return 7200
        elif self.displacement() <= 3600:
            return 8640
        elif self.displacement() <= 4200:
            return 9810
        elif self.displacement() <= 4800:
            return 11220
        else:
            return 12180


# 柴油車衍生類別
class DieselCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "柴油" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回燃料稅
    def fueltax(self):
        if self.displacement() <= 1800:
            return 2880
        elif self.displacement() <= 2400:
            return 3708
        elif self.displacement() <= 3000:
            return 4320
        elif self.displacement() <= 3600:
            return 5184
        elif self.displacement() <= 4200:
            return 5886
        elif self.displacement() <= 4800:
            return 6732
        else:
            return 7308


# 電能車衍生類別
class ElectricCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "電能" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回牌照稅
    def licensetax(self):
        return 0

    # 傳回燃料稅
    def fueltax(self):
        return 0


def tax(obj):
    print(obj.manufacture(), ":", obj.totaltax())


# 建構第1輛汽油車物件
print("==Test1==")
a = GasolineCars("CITROEN", "BX16TGS", "灰", 5, 1580)
print(type(a))
print(a.manufacture())
a.info()
print(a.fueltax())
print(a.licensetax())
print(a.totaltax())
# 可以修改屬性
a.color = "綠"
a.seat = 2
print(a.manufacture())
a.info()
# 建構第2輛柴油車物件
print("==Test2==")
b = DieselCars("福特六和", "C346-9W", "白", 5, 1997)
print(type(b))
print(b.manufacture())
b.info()
print(b.fueltax())
print(b.licensetax())
print(b.totaltax())
# 建構第3輛電能車物件
print("==Test3==")
c = ElectricCars("Tesla", "Model 3", "紅", 5, 346)
print(type(c))
print(c.manufacture())
c.info()
print(c.fueltax())
print(c.licensetax())
print(c.totaltax())
print("==Test4==")
tax(a)
tax(b)
tax(c)

print("------------------------------------------------------------")  # 60個


# chap8-02
# 輪胎類別
class Wheels:
    # 建構式
    def __init__(self, width, diameter):
        self.width = width  # 寬度屬性
        self.diameter = diameter  # 直徑屬性

    # 自訂義Wheels物件的比較關係(判斷是否相等)
    # 判斷兩個輪胎是否為相同類型根據輪胎寬度、直徑
    def __eq__(self, other):
        return self.width == other.width and self.diameter == other.diameter

    def eq(self, other):
        return self.width == other.width and self.diameter == other.diameter

    # 屬性(Attribute)
    def dimension(self):
        return "寬" + str(self.width) + "mm" + " 直徑" + str(self.diameter) + "inch"

    # 方法(Method)
    # 輸出輪胎屬性
    def info(self):
        print("寬" + str(self.width) + "mm" + " 直徑" + str(self.diameter) + "inch")


# 機車基礎類別
class Motorcycles:
    # 類別方法變數
    counts = 0

    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        self.__factory = factory  # 廠牌屬性
        self.__model = model  # 型號屬性
        self.color = color  # 顏色屬性
        self.__displacement = displacement  # 排氣量或電功率屬性
        self.wheel = wheel  # 輪胎屬性
        # 改變類別方法變數
        Motorcycles.counts = Motorcycles.counts + 1

    # 解構式
    def __del__(self):
        # 改變類別方法變數
        Motorcycles.counts = Motorcycles.counts - 1

    # 自訂義Motorcycles物件的比較關係(判斷是否相等)
    # 判斷兩台機車是否為相同類型根據廠商、型號、排氣量
    def __eq__(self, other):
        return (
            self.__factory == other.__factory
            and self.__model == other.__model
            and self.__displacement == other.__displacement
        )

    def eq(self, other):
        return (
            self.__factory == other.__factory
            and self.__model == other.__model
            and self.__displacement == other.__displacement
        )

    # 類別方法
    @classmethod
    # 顯示類別方法變數
    def motocycles_counts(cls):
        return cls.counts

    # 屬性(Attribute)
    # 函式唯讀
    # 傳回製造商出廠資料
    def manufacture(self):
        if self.__model.startswith("汽油"):
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "CC"
            )
        else:
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "HP"
            )

    # 傳回排氣量或電功率(保護)
    def displacement(self):
        return self.__displacement

    # 傳回燃料稅(以汽油機車為主，電能車不同，另外覆寫即可)
    def fueltax(self):
        if self.displacement() <= 50:
            return 300
        elif self.displacement() <= 125:
            return 450
        elif self.displacement() <= 250:
            return 600
        elif self.displacement() <= 500:
            return 900
        elif self.displacement() <= 600:
            return 1200
        elif self.displacement() <= 1200:
            return 1800
        else:
            return 2010

    # 傳回牌照稅(因為汽油是以排氣量計算，只有電能車不同，另外覆寫即可)
    def licensetax(self):
        if self.displacement() <= 150:
            return 0
        elif self.displacement() <= 250:
            return 800
        elif self.displacement() <= 500:
            return 1620
        elif self.displacement() <= 600:
            return 2160
        elif self.displacement() <= 1200:
            return 4320
        elif self.displacement() <= 1800:
            return 7120
        else:
            return 11230

    # 傳回全部稅金
    def totaltax(self):
        return self.fueltax() + self.licensetax()

    # 方法(Method)
    # 輸出機車屬性
    def info(self):
        if self.__model.startswith("汽油"):
            print(
                self.__factory
                + "-"
                + self.__model
                + "-"
                + self.color
                + "色-"
                + str(self.displacement())
                + "CC 輪胎"
                + self.wheel.dimension()
            )
        else:
            print(
                self.__factory
                + "-"
                + self.__model
                + "-"
                + self.color
                + "色-"
                + str(self.displacement())
                + "HP 輪胎"
                + self.wheel.dimension()
            )


class GasolineMotorcycles(Motorcycles):
    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        super().__init__(factory, "汽油" + model, color, displacement, wheel)


class ElectricMotorcycles(Motorcycles):
    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        super().__init__(factory, "電能" + model, color, displacement, wheel)

    # 傳回燃料稅(0元)
    def fueltax(self):
        return 0

    # 傳回牌照稅(0元)
    def licensetax(self):
        return 0


# 建立第1台機車
print("==Test1==")
d = GasolineMotorcycles("光陽", "SJ25JF", "白黑銀", 124, Wheels(100, 10))
# print(type(d))
print(d.manufacture())
d.info()
print(d.fueltax())
print(d.licensetax())
print(d.totaltax())
# 建立第2台機車
print("==Test2==")
e = GasolineMotorcycles("功學社", "CT-50", "白紅", 49, Wheels(100, 10))
# print(type(e))
print(e.manufacture())
e.info()
print(e.fueltax())
print(e.licensetax())
print(e.totaltax())
# 建立第3台機車
print("==Test3==")
f = ElectricMotorcycles("睿能", "GHS6B2", "白", 8.58, Wheels(100, 10))
# print(type(f))
print(f.manufacture())
f.info()
print(f.fueltax())
print(f.licensetax())
print(f.totaltax())
# 建立第4台機車
print("==Test4==")
g = GasolineMotorcycles("光陽", "SK60AF", "深藍", 298, Wheels(100, 10))
# print(type(g))
print(g.manufacture())
g.info()
print(g.fueltax())
print(g.licensetax())
print(g.totaltax())
# 建立第5台機車
print("==Test5==")
h = ElectricMotorcycles("睿能", "GHS6B2", "紅", 8.58, Wheels(100, 10))
# print(type(h))
print(h.manufacture())
h.info()
print(h.fueltax())
print(h.licensetax())
print(h.totaltax())
# 判斷f與h這兩台機車是否為相同類型
print("==Test6==")
f.info()
h.info()
print(f == h)
# 判斷d與e這兩台機車的輪台是否為相同類型
d.wheel.info()
e.wheel.info()
print(d.wheel.eq(e.wheel))
print(d.wheel == e.wheel)
print("==Test7==")
# 查看總共新增了幾台機車
print(Motorcycles.counts)
# 刪除d車
del d
print(Motorcycles.counts)
# 刪除e車
del e
print(Motorcycles.counts)
# 刪除f車
del f
print(Motorcycles.counts)
# 刪除g車
del g
print(Motorcycles.counts)
# 刪除h車
del h
print(Motorcycles.counts)

print("------------------------------------------------------------")  # 60個


# chap8-03
# 輪胎類別
class Wheels:
    # 建構式
    def __init__(self, width, diameter):
        self.width = width  # 寬度屬性
        self.diameter = diameter  # 直徑屬性

    # 自訂義Wheels物件的比較關係(判斷是否相等)
    # 判斷兩個輪胎是否為相同類型根據輪胎寬度、直徑
    def __eq__(self, other):
        return self.width == other.width and self.diameter == other.diameter

    def eq(self, other):
        return self.width == other.width and self.diameter == other.diameter

    # 屬性(Attribute)
    def dimension(self):
        return "寬" + str(self.width) + "mm" + " 直徑" + str(self.diameter) + "inch"

    # 方法(Method)
    # 輸出輪胎屬性
    def info(self):
        print("寬" + str(self.width) + "mm" + " 直徑" + str(self.diameter) + "inch")


# 車輛基礎類別
class Vehicles:
    # 建構式
    def __init__(self, factory, model, color, displacement):
        self.__factory = factory  # 廠牌屬性
        self.__model = model  # 型號屬性
        self.color = color  # 顏色屬性
        self.__displacement = displacement  # 排氣量或電功率屬性

    # 自訂義車輛物件的比較關係(判斷是否相等)
    # 判斷兩車輛是否為相同類型根據廠商、型號、排氣量
    def __eq__(self, other):
        return (
            self.__factory == other.__factory
            and self.__model == other.__model
            and self.__displacement == other.__displacement
        )

    def eq(self, other):
        return (
            self.__factory == other.__factory
            and self.__model == other.__model
            and self.__displacement == other.__displacement
        )

    # 屬性(Attribute)
    # 函式唯讀
    # 傳回製造商出廠資料
    def manufacture(self):
        if self.__model.startswith("汽油"):
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "CC"
            )
        else:
            return (
                self.__factory
                + "-"
                + self.__model
                + "-"
                + str(self.__displacement)
                + "HP"
            )

    # 傳回排氣量或電功率(保護)
    def displacement(self):
        return self.__displacement

    # 傳回燃料稅(另外覆寫即可)
    def fueltax(self):
        pass

    # 傳回牌照稅(因為汽油是以排氣量計算，只有電能車不同，另外覆寫即可)
    def licensetax(self):
        if self.displacement() <= 125:
            return 0
        elif self.displacement() <= 250:
            return 800
        elif self.displacement() <= 500:
            return 1620
        elif self.__displacement <= 600:
            return 2160
        elif self.__displacement <= 1200:
            return 4320
        elif self.__displacement <= 1800:
            return 7120
        elif self.__displacement <= 2400:
            return 11230
        elif self.__displacement <= 3000:
            return 15210
        elif self.__displacement <= 3600:
            return 28220
        elif self.__displacement <= 4200:
            return 28220
        elif self.__displacement <= 4800:
            return 46170
        else:
            return 46170

    # 傳回全部稅金
    def totaltax(self):
        return self.fueltax() + self.licensetax()

    # 方法(Method)
    # 輸出車輛屬性
    def info(self):
        if self.__model.startswith("電能"):
            if self.__model.find("機車") >= 0:
                print(
                    self.__factory
                    + "-"
                    + self.__model
                    + "-"
                    + self.color
                    + "色-"
                    + str(self.displacement())
                    + "HP 輪胎"
                    + self.wheel.dimension()
                )
            else:
                print(
                    self.__factory
                    + "-"
                    + self.__model
                    + "-"
                    + self.color
                    + "色-"
                    + str(self.seat)
                    + "人座"
                    + str(self.__displacement)
                    + "HP"
                )
        else:
            if self.__model.find("機車") >= 0:
                print(
                    self.__factory
                    + "-"
                    + self.__model
                    + "-"
                    + self.color
                    + "色-"
                    + str(self.displacement())
                    + "CC 輪胎"
                    + self.wheel.dimension()
                )
            else:
                print(
                    self.__factory
                    + "-"
                    + self.__model
                    + "-"
                    + self.color
                    + "色-"
                    + str(self.seat)
                    + "人座"
                    + str(self.__displacement)
                    + "CC"
                )


# 機車基礎類別
class Motorcycles(Vehicles):
    # 類別方法變數
    counts = 0

    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        super().__init__(factory, model, color, displacement)
        self.wheel = wheel  # 輪胎屬性
        # 改變類別方法變數
        Motorcycles.counts = Motorcycles.counts + 1

    # 解構式
    def __del__(self):
        # 改變類別方法變數
        Motorcycles.counts = Motorcycles.counts - 1

    # 類別方法
    @classmethod
    # 顯示類別方法變數
    def motocycles_counts(cls):
        return cls.counts

    # 屬性(Attribute)
    # 方法(Method)
    def fueltax(self):
        if self.displacement() <= 50:
            return 300
        elif self.displacement() <= 125:
            return 450
        elif self.displacement() <= 250:
            return 600
        elif self.displacement() <= 500:
            return 900
        elif self.displacement() <= 600:
            return 1200
        elif self.displacement() <= 1200:
            return 1800
        else:
            return 2010


class GasolineMotorcycles(Motorcycles):
    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        super().__init__(factory, "汽油機車" + model, color, displacement, wheel)


class ElectricMotorcycles(Motorcycles):
    # 建構式
    def __init__(self, factory, model, color, displacement, wheel):
        super().__init__(factory, "電能機車" + model, color, displacement, wheel)

    # 傳回燃料稅(0元)
    def fueltax(self):
        return 0

    # 傳回牌照稅(0元)
    def licensetax(self):
        return 0


# Cars基礎類別
class Cars(Vehicles):
    # 類別方法變數
    counts = 0

    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, model, color, displacement)
        self.seat = seat  # 座位屬性(可改)
        # 改變類別方法變數
        Cars.counts = Cars.counts + 1

    # 解構式
    def __del__(self):
        # 改變類別方法變數
        Cars.counts = Cars.counts - 1

    # 類別方法
    @classmethod
    # 顯示類別方法變數
    def cars_counts(cls):
        return cls.counts

    # 屬性(Attribute)
    # 方法(Method)


# 汽油車衍生類別
class GasolineCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "汽油汽車" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回燃料稅
    def fueltax(self):
        if self.displacement() <= 500:
            return 2160
        elif self.displacement() <= 600:
            return 2880
        elif self.displacement() <= 1200:
            return 4320
        elif self.displacement() <= 1800:
            return 4800
        elif self.displacement() <= 2400:
            return 6180
        elif self.displacement() <= 3000:
            return 7200
        elif self.displacement() <= 3600:
            return 8640
        elif self.displacement() <= 4200:
            return 9810
        elif self.displacement() <= 4800:
            return 11220
        else:
            return 12180


# 柴油車衍生類別
class DieselCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "柴油汽車" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回燃料稅
    def fueltax(self):
        if self.displacement() <= 1800:
            return 2880
        elif self.displacement() <= 2400:
            return 3708
        elif self.displacement() <= 3000:
            return 4320
        elif self.displacement() <= 3600:
            return 5184
        elif self.displacement() <= 4200:
            return 5886
        elif self.displacement() <= 4800:
            return 6732
        else:
            return 7308


# 電能車衍生類別
class ElectricCars(Cars):
    # 建構式
    def __init__(self, factory, model, color, seat, displacement):
        super().__init__(factory, "電能汽車" + model, color, seat, displacement)

    # 屬性(Attribute)
    # 傳回牌照稅(0元)
    def licensetax(self):
        return 0

    # 傳回燃料稅(0元)
    def fueltax(self):
        return 0


def tax(obj):
    print(obj.manufacture(), ":", obj.totaltax())


# 建構第1輛汽油車物件
print("==Test1==")
a = GasolineCars("CITROEN", "BX16TGS", "灰", 5, 1580)
# print(type(a))
print(a.manufacture())
a.info()
print(a.fueltax())
print(a.licensetax())
print(a.totaltax())
# 可以修改屬性
a.color = "綠"
a.seat = 2
print(a.manufacture())
a.info()
# 建構第2輛柴油車物件
print("==Test2==")
b = DieselCars("福特六和", "C346-9W", "白", 5, 1997)
# print(type(b))
print(b.manufacture())
b.info()
print(b.fueltax())
print(b.licensetax())
print(b.totaltax())
# 建構第3輛電能車物件
print("==Test3==")
c = ElectricCars("Tesla", "Model 3", "紅", 5, 346)
# print(type(c))
print(c.manufacture())
c.info()
print(c.fueltax())
print(c.licensetax())
print(c.totaltax())
print("==Test4==")
# 判斷a與b這兩輛汽車是否為相同類型
a.info()
b.info()
print(a == b)

# 建立第1台機車
print("==Test5==")
d = GasolineMotorcycles("光陽", "SJ25JF", "白黑銀", 124, Wheels(100, 10))
# print(type(d))
print(d.manufacture())
d.info()
print(d.fueltax())
print(d.licensetax())
print(d.totaltax())
# 建立第2台機車
print("==Test6==")
e = GasolineMotorcycles("功學社", "CT-50", "白紅", 49, Wheels(100, 10))
# print(type(e))
print(e.manufacture())
e.info()
print(e.fueltax())
print(e.licensetax())
print(e.totaltax())
# 建立第3台機車
print("==Test7==")
f = ElectricMotorcycles("睿能", "GHS6B2", "白", 8.58, Wheels(100, 10))
# print(type(f))
print(f.manufacture())
f.info()
print(f.fueltax())
print(f.licensetax())
print(f.totaltax())
# 建立第4台機車
print("==Test8==")
g = GasolineMotorcycles("光陽", "SK60AF", "深藍", 298, Wheels(100, 10))
# print(type(g))
print(g.manufacture())
g.info()
print(g.fueltax())
print(g.licensetax())
print(g.totaltax())
# 建立第5台機車
print("==Test9==")
h = ElectricMotorcycles("睿能", "GHS6B2", "紅", 8.58, Wheels(100, 10))
# print(type(h))
print(h.manufacture())
h.info()
print(h.fueltax())
print(h.licensetax())
print(h.totaltax())
print("==Test10==")
# 判斷f與h這兩台機車是否為相同類型
f.info()
h.info()
print(f == h)
print("==Test11==")
# 判斷d與e這兩台機車的輪台是否為相同類型
d.wheel.info()
e.wheel.info()
print(d.wheel.eq(e.wheel))
print(d.wheel == e.wheel)
print("==Test12==")
tax(a)
tax(b)
tax(c)
tax(d)
tax(e)
tax(f)
tax(g)
tax(h)
print("==Test13==")
# 查看總共新增了幾輛汽車
print(Cars.counts)
# 刪除a車
del a
print(Cars.counts)
# 刪除b車
del b
print(Cars.counts)
# 刪除c車
del c
print(Cars.counts)
print("==Test14==")
# 查看總共新增了幾台機車
print(Motorcycles.counts)
# 刪除d車
del d
print(Motorcycles.counts)
# 刪除e車
del e
print(Motorcycles.counts)
# 刪除f車
del f
print(Motorcycles.counts)
# 刪除g車
del g
print(Motorcycles.counts)
# 刪除h車
del h
print(Motorcycles.counts)

print("------------------------------------------------------------")  # 60個

# chap9-01b
import requests
import json

url = "https://od.cdc.gov.tw/eic/Day_Confirmation_Age_County_Gender_19CoV.json"
res = requests.get(url, verify=False)

print(res)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
print(data)

sys.exit()

# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 轉換'確定病例數'欄位內容為整數(以利後面加總)
df["確定病例數"] = df["確定病例數"].astype(int)

startdate = "2021/07/20"
# startdate=input("請輸日期(yyyy/mm/dd)")
# 只保留個案研判日、是否為境外移入、確定病例數 三個欄位
df.drop(["確定病名", "縣市", "鄉鎮", "性別", "年齡層"], inplace=True, axis=1)
# 設定過濾條件(累積)
indexNames = df[df["個案研判日"] > startdate].index
# 刪除所有大於指定日期的個案
df.drop(indexNames, inplace=True)
# 計算累積確診人數
total = df.sum()["確定病例數"]
print("累積確診人數:", total)
# 設定過濾條件(今日)
indexNames = df[df["個案研判日"] != startdate].index
# 刪除所有不是指定日期的個案
df.drop(indexNames, inplace=True)
# 計算境外移入人數
imported = df[df["是否為境外移入"] == "是"].sum()["確定病例數"]
print("今日境外移入人數:", imported)
# 計算本土人數
domestic = df[df["是否為境外移入"] == "否"].sum()["確定病例數"]
print("今日本土人數:", domestic)
total = imported + domestic
print("今日總人數:", total)
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)

print("------------------------------------------------------------")  # 60個

import requests

url = "https://www.dcard.tw/f/stock/p/237123381"
response = requests.get(url)
print(response.text)


print("chap7-02a")
import requests

url = "https://www.dcard.tw/f/stock/p/237123381"
res = requests.get(url)
print(res.text)

# chap7-01b
import requests

url = "http://jigsaw.w3.org/HTTP/connection.html"
response = requests.get(url)
# print(response.text)

# 在<HEAD></HEAD>區塊中取得包圍網頁標題的指定字串<TITLE></TITLE>所在的位置
# stripe()去除字串頭尾的'\n'(換行)、'\t'(跳格)、' '(空白)
datapos1 = response.text.find("<TITLE>")
datapos2 = response.text.find("</TITLE>")
data = response.text[datapos1 + 7 : datapos2]
data = data.strip()
print("網頁的<TITLE> :", data)
# 在<BODY></BODY>區塊中取得包圍內容標題的指定字串<H1></H1>所在的位置
datapos1 = response.text.find("<H1>")
datapos2 = response.text.find("</H1>")
data = response.text[datapos1 + 4 : datapos2]
# 將設定斜體的HTML語法<I></I>移除
data = data.replace("<I>", "")
data = data.replace("</I>", "")
data = data.strip()
print("<網頁的H1的資料(去掉<I>)> :", data)

datapos1 = response.text.find("<CODE>")
datapos2 = response.text.find("</CODE>")
data = response.text[datapos1 + 7 : datapos2]
data = data.strip()
print("網頁的<CODE> :", data)

print("------------------------------------------------------------")  # 60個

print("chap7-02b")
import requests
from bs4 import BeautifulSoup

url = "https://www.dcard.tw/f/stock/p/237123381"
response = requests.get(url)
# 指定html.parser作為解析器
soup = BeautifulSoup(response.text, "html.parser")
# 把排版後的html印出來，因為未排版前有很多網頁語法缺乏換行符號，不易閱讀
# 必須借助於Beautiful Shop套件
print(soup.prettify())

print("------------------------------------------------------------")  # 60個

print("chap7-02c")
import requests
from bs4 import BeautifulSoup

url = "https://www.dcard.tw/f/stock/p/237123381"

headers = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Accept-Language": "zh-TW,zh;q=0.9",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Upgrade-Insecure-Requests": "1",
    "User-Agent": "Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.102 Safari/537.36",  # 使用者代理
}

# response = requests.get(url="https://example.com", headers=headers)
response = requests.get(url, headers=headers)
# 指定html.parser作為解析器
soup = BeautifulSoup(response.text, "html.parser")
# 把排版後的html印出來
# print(soup.prettify())
a_tags = soup.find_all("h1")
print(">>>>>文章標題")
print(a_tags[0].contents[0])
print("\n")

print("------------------------------------------------------------")  # 60個

print("chap7-02d")

import requests
from bs4 import BeautifulSoup

url = "https://www.dcard.tw/f/stock/p/237123381"
res = requests.get(url)
# 指定html.parser作為解析器
soup = BeautifulSoup(res.text, "html.parser")
# 把排版後的html印出來
# print(soup.prettify())
a_tags = soup.find_all("h1")
print(">>>>>文章標題")
print(a_tags[0].contents[0])
print("\n")
a_tags = soup.find_all("div", limit=1)
a_tag = a_tags[0]
cc = ""
for b in a_tag.contents:
    if str(b).find("gFINpq") >= 0:
        b = str(b).replace("\n", "").replace("\r", "")
        cc = cc + b
print(cc.strip())

print("------------------------------------------------------------")  # 60個

print("chap7-02e")

import requests
from bs4 import BeautifulSoup

url = "https://www.dcard.tw/f/stock/p/237123381"
res = requests.get(url)
# 指定html.parser作為解析器
soup = BeautifulSoup(res.text, "html.parser")
# 把排版後的html印出來
# print(soup.prettify())
a_tags = soup.find_all("h1")
print(">>>>>文章標題")
print(a_tags[0].contents[0])
print()
a_tags = soup.find_all("div", limit=1)
a_tag = a_tags[0]
cc = ""
for b in a_tag.contents:
    if str(b).find("gFINpq") >= 0:
        b = str(b)
        cc = cc + b
cc = cc.strip()
data = ">>>>>原Po文章\n"
# 尋找前四筆，第一個是原Po文章，後三個則為熱門留言
for j in range(4):
    # 尋找最附近的<div class>有gFINpq
    datapos1 = cc.find("gFINpq")
    # 只保留這個字串以後的文字
    cc = cc[datapos1:]
    while True:
        # 尋找最附近的<span
        datapos1 = cc.find("<span")
        # 只保留<span以後的文字
        cc = cc[datapos1:]
        # 尋找最附近的>
        datapos1 = cc.find(">")
        # 尋找最附近的</span>
        datapos2 = cc.find("</span>")
        # 如果<span></span>中間出現了enUbOQ，表示這個步驟該結束了
        if cc[:datapos2].find("enUbOQ") >= 0:
            break
        # 如果<span></span>有資料，就合併在data字串裡
        if datapos1 + 1 < datapos2:
            # 但是<span></span>的資料，如果還有span就不行了
            if cc[datapos1 + 1 : datapos2].find("span") < 0:
                data = data + cc[datapos1 + 1 : datapos2] + "\n"
        # 只保留</span>之後的文字
        cc = cc[datapos2 + 7 :]
    # 後三個是熱門留言
    if j < 3:
        data = data + "\n*******熱門留言" + str(j + 1) + ":\n"
print(data)

print("------------------------------------------------------------")  # 60個

print("chap7-02f")

import requests
import json

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 查看有哪些欄位
print(data.keys())
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 查看有哪些內容
print(df)

print("------------------------------------------------------------")  # 60個

print("chap7-02g")

import requests
import json

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df.loc["title", 0])
print()
print(">>>>>原Po文章")
print(str(df.loc["content", 0]).strip())
print("------------------------------------------------------------")  # 60個

print("chap7-02h")

import requests
import json
from datetime import datetime

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df[0]["title"])
print()
print(">>>>>原Po文章\n")
print(str(df[0]["content"]).strip())
url = "https://www.dcard.tw/service/api/v2/posts/" + postid + "/comments?popular=True"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 取出前3筆資料
for i in range(3):
    print("*******熱門留言" + str(i) + ":")
    print(df.loc[i, "content"])
    # 將最後修改文字日期轉換成日期
    updatedate = datetime.fromisoformat(str(df.loc[i, "updatedAt"]).replace("Z", ""))
    # 計算現在和最後修改日期的時間差
    datediff = datetime.today() - updatedate
    # 顯示幾天前有最新留言
    print(datediff.days, "天前")
    print()
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)

print("------------------------------------------------------------")  # 60個

print("chap7-02i")

import requests
import json
from datetime import datetime

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df[0]["title"])
print()
print(">>>>>原Po文章\n")
print(str(df[0]["content"]).strip())
url = "https://www.dcard.tw/service/api/v2/posts/" + postid + "/comments"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 除了updatedAt, content, likeCount三個欄位以外，全部刪除
# 刪除欄依定要指定axis=1
# inplace=True真實刪除
df.drop(
    [
        "id",
        "anonymous",
        "postId",
        "createdAt",
        "floor",
        "withNickname",
        "hiddenByAuthor",
        "meta",
        "gender",
        "school",
        "host",
        "reportReason",
        "mediaMeta",
        "hidden",
        "inReview",
        "reportReasonText",
        "isSuspiciousAccount",
        "isModerator",
        "doorplate",
        "edited",
        "postAvatar",
        "activityAvatar",
        "verifiedBadge",
        "memberType",
        "enablePrivateMessage",
        "department",
    ],
    inplace=True,
    axis=1,
)
# 依據likeCout內容排序，以降序排列
df.sort_values(by="likeCount", inplace=True, ascending=False)
# 要記得重設index，這樣df.loc[]的結果才會正確
df = df.reset_index(drop=True)
# 取出前3筆資料
for i in range(3):
    print("*******熱門留言" + str(i) + ":")
    print(df.loc[i, "content"])
    # 將最後修改文字日期轉換成日期
    updatedate = datetime.fromisoformat(str(df.loc[i, "updatedAt"]).replace("Z", ""))
    # 計算現在和最後修改日期的時間差
    datediff = datetime.today() - updatedate
    # 顯示幾天前有最新留言
    print(datediff.days, "天前")
    print()
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)

print("------------------------------------------------------------")  # 60個

print("chap7-02j")

import requests
import csv

# from io import StringIO

county = "屏東縣"
url = "https://data.epa.gov.tw/api/v1/aqx_p_432?limit=1000&api_key=keykeykey&sort=ImportDate%20desc&format=csv"
# 方法1:下載檔案後儲存，再開啟檔案讀出
# res = requests.get(url)
# open('data.csv','wb').write(res.content)
# df=pd.read_csv('data.csv')
# 方法2:直接讀取網站的檔案
df = pd.read_csv(url)
# 查看有那些欄位
df.info()
# 只保留SiteName、County、AQI、Status四個欄位
df.drop(
    [
        "SiteId",
        "Longitude",
        "Latitude",
        "PublishTime",
        "SO2_AVG",
        "Pollutant",
        "SO2",
        "CO",
        "NO",
        "CO_8hr",
        "O3",
        "O3_8hr",
        "PM10",
        "PM10_AVG",
        "PM2.5",
        "PM2.5_AVG",
        "NO2",
        "NOx",
        "WindSpeed",
        "WindDirec",
    ],
    inplace=True,
    axis=1,
)
# 設定過濾條件
indexNames = df[df["County"] != county].index
# 刪除所有不是臺北市的偵測站
df.drop(indexNames, inplace=True)
# 重設索引值
df = df.reset_index(drop=True)
# 列出該縣市所有偵測站狀態
print(county, "空氣品質狀態")
for i in range(len(df)):
    print(df.loc[i, "SiteName"], "(", df.loc[i, "Status"], ")")
indexNames = df[df["Status"] == "良好"].index
df.drop(indexNames, inplace=True)
# 重設索引值
df = df.reset_index(drop=True)
# 列出該縣市所有為達良好之偵測站
print()
print(county, "空氣品質未達良好之偵測站")
for i in range(len(df)):
    print(df.loc[i, "SiteName"], "(", df.loc[i, "Status"], ")")

print("------------------------------------------------------------")  # 60個

# chap9-01b
import requests
import json

url = "https://od.cdc.gov.tw/eic/Day_Confirmation_Age_County_Gender_19CoV.json"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 轉換'確定病例數'欄位內容為整數(以利後面加總)
df["確定病例數"] = df["確定病例數"].astype(int)

startdate = "2021/07/20"
# startdate=input("請輸日期(yyyy/mm/dd)")
# 只保留個案研判日、是否為境外移入、確定病例數 三個欄位
df.drop(["確定病名", "縣市", "鄉鎮", "性別", "年齡層"], inplace=True, axis=1)
# 設定過濾條件(累積)
indexNames = df[df["個案研判日"] > startdate].index
# 刪除所有大於指定日期的個案
df.drop(indexNames, inplace=True)
# 計算累積確診人數
total = df.sum()["確定病例數"]
print("累積確診人數:", total)
# 設定過濾條件(今日)
indexNames = df[df["個案研判日"] != startdate].index
# 刪除所有不是指定日期的個案
df.drop(indexNames, inplace=True)
# 計算境外移入人數
imported = df[df["是否為境外移入"] == "是"].sum()["確定病例數"]
print("今日境外移入人數:", imported)
# 計算本土人數
domestic = df[df["是否為境外移入"] == "否"].sum()["確定病例數"]
print("今日本土人數:", domestic)
total = imported + domestic
print("今日總人數:", total)
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)

print("------------------------------------------------------------")  # 60個

# chap9-01c
import requests
import csv

# from io import StringIO

url = "https://od.cdc.gov.tw/eic/covid19/covid19_tw_stats.csv"
# 方法1:下載檔案後儲存，再開啟檔案讀出
# res = requests.get(url)
# open('data.csv','wb').write(res.content)
# df=pd.read_csv('data.csv')
# 方法2:直接讀取網站的檔案
df = pd.read_csv(url)
# 查看有那些欄位
df.info()
# 只保留確診、死亡、送驗、排除四個欄位
df.drop(["昨日確診", "昨日排除", "昨日送驗"], inplace=True, axis=1)
print(df)
totaldeath = df.loc[0, "死亡"]
print("累積死亡人數", totaldeath)

print("------------------------------------------------------------")  # 60個

# chap9-01d
import requests
import json
import csv
from docx import Document

url = "https://od.cdc.gov.tw/eic/Day_Confirmation_Age_County_Gender_19CoV.json"
res = requests.get(url)
data = res.json()
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 轉換'確定病例數'欄位內容為整數(以利後面加總)
df["確定病例數"] = df["確定病例數"].astype(int)
startdate = "2021/07/20"
startdate = input("請輸日期(yyyy/mm/dd)")
# 只保留個案研判日、是否為境外移入、確定病例數 三個欄位
df.drop(["確定病名", "縣市", "鄉鎮", "性別", "年齡層"], inplace=True, axis=1)
# 設定過濾條件(累積)
indexNames = df[df["個案研判日"] > startdate].index
# 刪除所有大於指定日期的個案
df.drop(indexNames, inplace=True)
# 計算累積確診人數
total = df.sum()["確定病例數"]
print("累積確診人數:", total)
# 設定過濾條件(今日)
indexNames = df[df["個案研判日"] != startdate].index
# 刪除所有不是指定日期的個案
df.drop(indexNames, inplace=True)
# 計算境外移入人數
imported = df[df["是否為境外移入"] == "是"].sum()["確定病例數"]
print("今日境外移入人數:", imported)
# 計算本土人數
domestic = df[df["是否為境外移入"] == "否"].sum()["確定病例數"]
print("今日本土人數:", domestic)
print("今日總人數:", domestic + imported)

url = "https://od.cdc.gov.tw/eic/covid19/covid19_tw_stats.csv"
df = pd.read_csv(url)
# 只保留確診、死亡、送驗、排除四個欄位
df.drop(["昨日確診", "昨日排除", "昨日送驗"], inplace=True, axis=1)
totaldeath = df.loc[0, "死亡"]
print("累積死亡人數", totaldeath)

# 切換至指定的資料夾為工作資料夾
os.chdir("/content/drive/MyDrive/Book")
# 開啟指定的Word檔案
my_doc = Document("指揮中心快訊.docx")
# 設定要取代的串列list
replacements = {
    "%Date%": startdate,
    "%N1%": str(domestic + imported),
    "%N2%": str(domestic),
    "%N3%": str(imported),
    "%N4%": str(totaldeath),
    "%N5%": str(total),
}
# 尋找要取代的關鍵字
for key in replacements:
    # 尋找所有的表格
    for table in my_doc.tables:
        # 尋找表格中的列
        for row in table.rows:
            # 尋找列中的每個儲存格
            for cell in row.cells:
                # 尋找儲存格中的每一個斷落
                for paragraph in cell.paragraphs:
                    # 如果關鍵字出現在段落中
                    if key in paragraph.text:
                        # 為了避免修改掉原有的樣式，必須用這個方法處理
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                text = inline[i].text.replace(key, replacements[key])
                                inline[i].text = text
# 指定儲存的檔案名稱
fname = "指揮中心快訊" + startdate.replace("/", "") + ".docx"
# 儲存檔案
my_doc.save(fname)

"""
請輸日期(yyyy/mm/dd)2021/07/20

累積確診人數: 15450

今日境外移入人數: 6

今日本土人數: 18

今日總人數: 24

累積死亡人數 846
"""

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

try:
    # 嘗試打開一個不存在的檔案
    with open("non_existent_file.txt", "r") as f:
        data = f.read()  # read(無參數), 從目前檔案位置讀到檔尾
except FileNotFoundError:
    # 如果文件不存在, 捕獲異常
    print("The file was not found")
except IOError:
    # 處理 I/O 錯誤, 例如:讀取錯誤
    print("An I/O error occurred")

print("------------------------------------------------------------")  # 60個

filename = "data15_4.txt"
try:
    with open(filename) as f:  # 預設mode=r開啟檔案
        data = f.read()  # read(無參數), 從目前檔案位置讀到檔尾
except FileNotFoundError:
    print(f"找不到 {filename} 檔案")
else:
    print(data)  # 輸出變數data


print("------------------------------------------------------------")  # 60個


set99 = set()
for i in range(2, 9 + 1):
    set99.add(i)

print(type(set99))
print(set99)


print("------------------------------------------------------------")  # 60個

import os

pName = "C:/pcYah"
if os.path.isdir(pName):  # 檢查資料夾路徑是否存在
    print("%s 資料夾路徑存在" % pName)
else:
    print("%s 資料夾路徑不存在" % pName)

fName = "C:/Windows/win.ini"
if os.path.isfile(fName):  # 檢查檔案路徑是否存在
    print("%s 檔案路徑存在" % fName)
else:
    print("%s 檔案路徑不存在" % fName)

"""

isdir  isfile

if os.path.exists(pName):        # 檢查資料夾路徑是否存在
if os.path.exists(fName):        # 檢查檔案路徑是否存在


"""


# try-catch-finally
n1 = 8
n2 = 0
try:
    d = n1 / n2
    print("%d / %d = %d" % (n1, n2, d))
except Exception as e:
    print("錯誤類型 :", end=" ")
    print(e)
finally:
    print("執行 finally: 敘述\n")


print("------------------------------------------------------------")  # 60個


lst = [0 for x in range(4)]
try:
    lst[3] = 33
    print("lst[3] =", lst[3])
    lst[8] = 88
    print("lst[8] =", lst[8])
except ZeroDivisionError:
    print("錯誤類型 : 除數為零")
except IndexError:
    print("錯誤類型 : 串列註標超出範圍")
except MemoryErroe:
    print("錯誤類型 : 超出記憶體空間")
except Exception as e:
    print("錯誤類型 :", e)

import os

fName = "score.txt"
if os.path.isfile(fName):
    f = open(fName, "r")
    print("讀1行")
    str1 = f.readline()
    print(str1)
    print("讀4行")
    str2 = f.readline(4)
    print(str2)
    print("剩下的讀完")
    print(f.read())
    f.close()
else:
    print(None)


print("------------------------------------------------------------")  # 60個

# ord(x) 可以將參數x所代表的Unicode字元，轉換為對應編碼數字
A = ord("A")
B = ord("B")
C = ord("C")

print(A)
print(B)
print(C)

# chr(x) 可以將參數x轉換為所代表的Unicode字元
AA = chr(A)
BB = chr(B)
CC = chr(C)

print(AA)
print(BB)
print(CC)

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個

""" fail
import time as t

timestart = t.clock()
for i in range (0,1000):
    for j in range (0,1000):
        n = i * j
timeend = t.clock()
print("執行一百萬次整數運算的時間：" + str(timeend-timestart) + " 秒")
"""

print("------------------------------------------------------------")  # 60個

import time as t

week = ["一", "二", "三", "四", "五", "六", "日"]
dst = ["無日光節約時間", "有日光節約時間"]
time1 = t.localtime()
show = "現在時刻：中華民國 " + str(int(time1.tm_year) - 1911) + " 年 "
show += str(time1.tm_mon) + " 月 " + str(time1.tm_mday) + " 日 "
show += str(time1.tm_hour) + " 點 " + str(time1.tm_min) + " 分 "
show += str(time1.tm_sec) + " 秒 星期" + week[time1.tm_wday] + "\n"
show += "今天是今年的第 " + str(time1.tm_yday) + " 天，此地" + dst[time1.tm_isdst]
print(show)

print("------------------------------------------------------------")  # 60個

import time as t

timestart = t.perf_counter()
for i in range(0, 1000):
    for j in range(0, 1000):
        n = i * j
timeend = t.perf_counter()
print("執行一百萬次整數運算的時間：" + str(timeend - timestart) + " 秒")

print("------------------------------------------------------------")  # 60個

"""
import time as t

timestart = t.clock()
for i in range (0,1000):
    for j in range (0,1000):
        n = float(i) * float(j)
timeend = t.clock()
print("執行一百萬次浮點數運算的時間：" + str(timeend-timestart) + " 秒")
"""
print("------------------------------------------------------------")  # 60個

date1 = "2017-8-23"
date1 = "西元 " + date1
date1 = date1.replace("-", " 年 ", 1)
date1 = date1.replace("-", " 月 ", 1)
date1 += " 日"
print(date1)

print("------------------------------------------------------------")  # 60個

import time as t

time1 = t.localtime()
show = "今天是今年的第 " + str(time1.tm_yday) + " 天，屬於"
if time1.tm_yday < 184:
    show += "上半年。"
else:
    show += "下半年。"
print(show)

print("------------------------------------------------------------")  # 60個

"""
import time as t

timestart = t.perf_counter() ()
for i in range (0,1000):
    for j in range (0,1000):
        n = float(i) * float(j)
timeend = t.perf_counter() ()
print("執行一百萬次浮點數運算的時間：" + str(timeend-timestart) + " 秒")
"""
print("------------------------------------------------------------")  # 60個

time1 = "10:23:41"
time1 = time1.replace(":", " 點 ", 1)
time1 = time1.replace(":", " 分 ", 1)
time1 += " 秒"
print(time1)

print("------------------------------------------------------------")  # 60個

import random

list1 = random.sample(range(1, 50), 7)
special = list1.pop()
list1.sort()
print("本期大樂透中獎號碼為：", end="")
for i in range(0, 6):
    if i == 5:
        print(str(list1[i]))
    else:
        print(str(list1[i]), end=", ")
print("本期大樂透特別號為：" + str(special))

print("------------------------------------------------------------")  # 60個

import time as t

time1 = t.localtime()
show = "現在時刻："
if time1.tm_hour < 12:
    show += "上午 "
    hour = time1.tm_hour
else:
    show += "下午 "
    hour = time1.tm_hour - 12
show += str(hour) + " 點 " + str(time1.tm_min) + " 分 "
show += str(time1.tm_sec) + " 秒"
print(show)

print("------------------------------------------------------------")  # 60個

dict1 = {"A": "內向穩重", "B": "外向樂觀", "O": "堅強自信", "AB": "聰明自然"}
name = "O"
blood = dict1.get(name)
if blood == None:
    print("沒有「" + name + "」血型！")
else:
    print(name + " 血型的個性為：" + str(dict1[name]))

print("------------------------------------------------------------")  # 60個

person = 7
apple = 52
ret = divmod(apple, person)
print("每個學生可分得蘋果 " + str(ret[0]) + " 個")
print("蘋果剩餘 " + str(ret[1]) + " 個")

print("------------------------------------------------------------")  # 60個

listname = ["林大明", "陳阿中", "張小英"]
listchinese = [100, 74, 82]
listmath = [87, 88, 65]
listenglish = [79, 100, 8]
print("姓名     座號  國文  數學  英文")
for i in range(0, 3):
    print(
        listname[i].ljust(5),
        str(i + 1).rjust(3),
        str(listchinese[i]).rjust(5),
        str(listmath[i]).rjust(5),
        str(listenglish[i]).rjust(5),
    )

print("------------------------------------------------------------")  # 60個


def disp_data():  # 顯示串列的自訂程序
    for item in datas:
        print(item, end=" ")
    print()


# 主程式
datas = [3, 5, 2, 1]
print("排序前：", end=" ")
disp_data()  # 顯示排序前串列
n = len(datas) - 1  # 串列長度-1

for i in range(0, n):
    for j in range(0, n - i):
        if datas[j] > datas[j + 1]:  # 由小到大排序
            datas[j], datas[j + 1] = datas[j + 1], datas[j]  # 兩數互換

print("排序後：", end=" ")
disp_data()  # 顯示排序後串列

print("------------------------------------------------------------")  # 60個


def disp_data():  # 顯示串列的自訂程序
    for item in datas:
        print(item, end=" ")
    print()


# 主程式
datas = [3, 5, 2, 1]
print("排序前：", end=" ")
disp_data()  # 顯示排序前串列
n = len(datas) - 1  # 串列長度-1

for i in range(0, n):
    for j in range(0, n - i):
        print("i=%d j=%d" % (i, j))
        if datas[j] > datas[j + 1]:  # 由大到小排序
            print("%d,%d 互換後" % (datas[j], datas[j + 1]), end="：")
            datas[j], datas[j + 1] = datas[j + 1], datas[j]  # 兩數互換
        print(datas)

print("排序後：", end=" ")
disp_data()  # 顯示排序後串列

print("------------------------------------------------------------")  # 60個

dict1 = {"春季": "暖和", "夏季": "炎熱", "秋季": "涼爽", "冬季": "寒冷"}
name = "春季"
feeling = dict1.get(name)
if feeling == None:
    print("沒有「" + name + "」季節！")
else:
    print(name + "的天氣為 " + str(dict1[name]))

print("------------------------------------------------------------")  # 60個

dict1 = {"電視": 15000, "冰箱": 23000, "冷氣": 28000}
name = "冰箱"
if name in dict1:
    print(name + "的價格為 " + str(dict1[name]))
else:
    price = input("輸入電器價格：")
    dict1[name] = price
    print("字典內容：" + str(dict1))

print("------------------------------------------------------------")  # 60個

dict1 = {"水瓶座": "活潑善變", "雙魚座": "迷人保守", "白羊座": "天生勇者", "金牛座": "熱情敏感"}
item1 = dict1.items()
for name, nature in item1:
    print("%s 的性格特癥為 %s" % (name, nature))

print("------------------------------------------------------------")  # 60個

dict1 = {"水瓶座": "活潑善變", "雙魚座": "迷人保守", "白羊座": "天生勇者", "金牛座": "熱情敏感"}
listkey = list(dict1.keys())
listvalue = list(dict1.values())
for i in range(len(listkey)):
    print("%s 的性格特癥為 %s" % (listkey[i], listvalue[i]))

print("------------------------------------------------------------")  # 60個

ret = divmod(10000, 350)
print("可維持生活 " + str(ret[0]) + " 天")
print("生活費剩餘 " + str(ret[1]) + " 元")

print("------------------------------------------------------------")  # 60個

listname = ["鍾明達", "鄭廣月", "何美麗"]
list1 = [34210, 23600, 145000, 54300]
list2 = [9000, 23900, 83400, 132000]
list3 = [186500, 127800, 100000, 45000]
list4 = [78900, 125000, 90000, 8000]
print("姓名       第一季  第二季  第三季   第四季")
for i in range(0, 3):
    print(
        listname[i].ljust(5),
        str(list1[i]).rjust(7),
        str(list2[i]).rjust(7),
        str(list3[i]).rjust(7),
        str(list4[i]).rjust(7),
    )

print("------------------------------------------------------------")  # 60個

import random

list1 = random.sample(range(0, 10), 4)
list1.sort()
print("本期四星彩中獎號碼為：", end="")
for i in range(0, 4):
    if i == 3:
        print(str(list1[i]))
    else:
        print(str(list1[i]), end=", ")

print("------------------------------------------------------------")  # 60個


def disp_data():  # 顯示串列的自訂程序
    for item in datas:
        print(item, end=" ")
    print()


# 主程式
datas = datas = [2, 3, 5, 7, 1]
print("排序前：", end=" ")
disp_data()  # 顯示排序前串列
n = len(datas) - 1  # 串列長度-1

for i in range(0, n):
    for j in range(0, n - i):
        if datas[j] < datas[j + 1]:  # 由大到小排序
            datas[j], datas[j + 1] = datas[j + 1], datas[j]  # 兩數互換

print("排序後：", end=" ")
disp_data()  # 顯示排序後串列

print("------------------------------------------------------------")  # 60個

monthname = {
    1: "JAN",
    2: "FEB",
    3: "MAR",
    4: "APR",
    5: "MAY",
    6: "JUN",
    7: "JUL",
    8: "AUG",
    9: "SEP",
    10: "OCT",
    11: "NOV",
    12: "DEC",
}
m = 3
print("{}月份的英文簡寫為 {}：".format(m, monthname[m]))

print("------------------------------------------------------------")  # 60個

cnum = {0: "零", 1: "壹", 2: "貮", 3: "參", 4: "肆", 5: "伍", 6: "陸", 7: "柒", 8: "捌", 9: "玖"}
num = "1234"
for n in num:
    print(cnum[int(n)], end="")

print("------------------------------------------------------------")  # 60個

dict1 = {"台北市": 6, "新北市": 2, "桃園市": 5, "台中市": 8, "台南市": 3, "高雄市": 9}
name = "新北市"
PM25 = dict1.get(name)
if PM25 == None:
    print("六都中沒有「" + name + "」城市！")
else:
    print(name + " 今天的 PM2.5 值為：" + str(dict1[name]))

print("------------------------------------------------------------")  # 60個

dict1 = {"鼠": "親切和藹", "牛": "保守努力", "虎": "熱情大膽", "兔": "溫柔仁慈"}
for name, nature in dict1.items():
    print("生肖屬 %s 的性格特癥為 %s" % (name, nature))

print("------------------------------------------------------------")  # 60個

rate = {"USD": 28.02, "JPY": 0.2513, "CNY": 4.24}
TWD = float("123.456")
print(
    "台幣{:.2f}元等於美金{:.2f}元, 日幣{:.2f}元, 人民幣{:.2f}元".format(
        TWD, TWD / rate["USD"], TWD / rate["JPY"], TWD / rate["CNY"]
    )
)

print("------------------------------------------------------------")  # 60個

import random

print("產生N個 從 MIN 到 MAX 不重複的整數(包含頭尾)")
N = 7
MIN = 1
MAX = 50
list1 = random.sample(range(MIN, MAX), N)
print(type(list1))
print(list1)
list1 = random.sample(range(MIN, MAX), N)
print(list1)
list1 = random.sample(range(MIN, MAX), N)
print(list1)
list1 = random.sample(range(MIN, MAX), N)
print(list1)
list1 = random.sample(range(MIN, MAX), N)
print(list1)

print("------------------------------------------------------------")  # 60個

N = 10
MIN = 80
MAX = 100
scores = random.sample(range(MIN, MAX), N)
print("原成績：", scores)

print("人數：%d" % len(scores))
print("最高分為：%d" % max(scores))
print("最低分為：%d" % min(scores))
print("總分為：%d" % sum(scores))
print("平均為：%6.2f" % (sum(scores) / N))

scores2 = sorted(scores, reverse=True)  # 由大到小排序
print("成績由大到小排序：", scores2)

scores2 = sorted(scores, reverse=False)  # 由小到大排序
print("成績由小到大排序：", scores2)

print("------------------------------------------------------------")  # 60個

N = 10
MIN = 80
MAX = 100
scores = random.sample(range(MIN, MAX), N)


def disp_scores():  # 顯示串列的自訂程序
    for score in scores:
        print(score, end=" ")
    print()


print("排序前：", end=" ")
disp_scores()  # 顯示排序前串列

n = len(scores) - 1  # 串列長度-1

for i in range(0, n):
    for j in range(0, n - i):
        if scores[j] < scores[j + 1]:  # 由大到小排序
            scores[j], scores[j + 1] = scores[j + 1], scores[j]  # 兩數互換

print("成績由大到小排序：", end="")
disp_scores()  # 顯示排序後串列

print("------------------------------------------------------------")  # 60個


sys.exit()

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個

animals = ["鼠", "牛", "虎", "兔", "龍"]

print("動物有：", animals)

animal = "豬"
n = animals.count(animal)
if n > 0:  # 串列元素存在
    p = animals.index(animal)
    print("%s 在串列中的第 %d 項" % (animal, p + 1))
    animals.remove(animal)
else:
    print(animal, "不在串列中!, 加入此動物")
    animals.append(animal)

print("動物有：", animals)

print("------------------------------------------------------------")  # 60個

dict1 = {"金牌": 26, "銀牌": 34, "銅牌": 30}
listkey = list(dict1.keys())
listvalue = list(dict1.values())
for i in range(len(listkey)):
    print("得到的 %s 數目為 %d 面" % (listkey[i], listvalue[i]))


print("------------------------------------------------------------")  # 60個

names = ["David", "Lily", "Chiou", "Bear", "Shantel", "Cynthia"]

n = len(names) - 1  # 串列長度-1
for i in range(0, n):
    for j in range(0, n - i):
        if names[j] > names[j + 1]:  # 由小到大排序
            names[j], names[j + 1] = names[j + 1], names[j]  # 互換
print("------------------------------------------------------------")  # 60個

print("姓名    成績")
print("%-4s  %3d" % (name1, score1))
print("%-4s  %3d" % (name2, score2))

print("------------------------------------------------------------")  # 60個

import sys
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

print("------------------------------------------------------------")  # 60個

# 過度擬合 (overfitting)

# 拉格朗日 (Lagrange) 插值法

x = np.linspace(0, 1, 200)
y = -((x - 1) ** 2) + 1
plt.plot(x, y, "lime")

X = np.linspace(0, 1, 20)
Y = -((X - 1) ** 2) + 1 + 0.08 * np.random.randn(20)
plt.scatter(X, Y, c="b", s=50)

z = np.polyfit(X, Y, 19)
p = np.poly1d(z)
plt.plot(x, p(x), "r")

xmin, xmax, ymin, ymax = 0, 1, 0, 1.5
plt.axis([xmin, xmax, ymin, ymax])  # 設定各軸顯示範圍

plt.show()

print("------------------------------------------------------------")  # 60個

"""
#【秘技】分列 X, Y 的變成點座標

等一下我們會大量的把資料變換形式, 現在我們先熱身。在畫圖時常常用到把 x, y 座標分列。現在我們要合成點要怎麼做呢? 也就是說
X1 <-- [1,2,3,4]
Y1 <-- [5,6,7,8]
希望變成
[[1,5], [2,6], [3,7], [4,8]]
"""

X1 = np.array([1, 2, 3, 4])
Y1 = np.array([5, 6, 7, 8])

# NumPy 有個神奇的方式會幫我們做!

ccc = np.c_[X1, Y1]
print(ccc)

"""
【重要插播】meshgrid 用法

為了用 contourf (填充型的等高線) 呈現我們成果, 我們要介紹一個初學有點難理解、meshgrid 的概念。

meshgrid 是產生格點的方式, 通常是我們要畫 3D 曲面啦、或是等高線的時候要先為我們在 xy 平面上「佈點」, 然後算出每點的高度 Z。

我們要做的是給定 x 方向座標, y 方向座標, 然後就產生格點, 如圖示。
"""


# 於是我們再度用我們的 X1, Y1 示範。

X1 = np.array([1, 2, 3, 4])

Y1 = np.array([5, 6, 7, 8])

# 因為 matplotlib 很愛 x, y-座標分開, 經 meshgrid 後也是分開的! 所以我們用 Xm 和 Ym 來接。

Xm, Ym = np.meshgrid(X1, Y1)

# 看一下內容...

print(Xm)
print(Ym)

# 等等, 這什麼啊? 原來 meshgrid 存網格的 x 座標是一列一列存的。
# 同理我們可以理解 Ym 的內容為什麼是這樣了...

print("------------------------------------------------------------")  # 60個

# 01 numpy 的 filter

egg = np.array([3, -5, 10, 23, -5, 11])
idx = egg >= 0
print(idx)
# array([ True, False,  True,  True, False,  True])

print(egg[idx])
# array([ 3, 10, 23, 11])

print(egg[egg >= 0])
# array([ 3, 10, 23, 11])

x = np.linspace(-10, 10, 1000)
y = np.sin(x)
plt.plot(x, y)
plt.scatter(x[y > 0], y[y > 0], c="r")
plt.show()

# 02 Overfitting
Px = np.random.rand(6)
Py = np.random.rand(6)
plt.scatter(Px, Py, c="r", s=50)
plt.show()

x = np.linspace(0, 1, 1000)
y = 0.5 * np.sin(x) + 0.5
plt.scatter(Px, Py, c="r", s=50)
plt.plot(x, y)
plt.show()


def myplot(n=1):
    y = 0.5 * np.sin(n * x) + 0.5
    plt.scatter(Px, Py, c="r", s=50)
    plt.plot(x, y)


myplot(3)

plt.show()

print("------------------------------------------------------------")  # 60個


# lambda: 臨時要使用的函數
currency = 32.1357851  # 1美元 = 32.13台幣    台灣銀行 現金賣出價

price = [100, 500, 1000]  # 美元

ll = list(map(lambda x: currency * x, price))
print("換算成台幣 :", ll)

usb2twd = lambda x: currency * x

print(usb2twd(1000))


# print(f"1 美元合台幣 {c:.2f} 元。")
# print(f"1 美元合台幣 {c:10.2f} 元。")

print("------------------------------------------------------------")  # 60個


class Card:
    SUITS = ["♣", "♦", "♥", "♠"]
    RANKS = ["A", "2", "3", "4", "5", "6", "7", "8", "9", "10", "J", "Q", "K"]

    def __init__(self, s, r):
        self.suit = s
        self.rank = r

    def show(self):
        print(self.SUITS[self.suit] + self.RANKS[self.rank])


card01 = Card(2, 3)
card01.show()

print("------------------------------------------------------------")  # 60個

"""
import matplotlib as mpl
import matplotlib.font_manager as fm

for f in mpl.font_manager.fontManager.ttflist:
    print(f.name)

#[f.name for f in mpl.font_manager.fontManager.ttflist]

print('matplotlib 真的「看到的」字型')
for f in fm.fontManager.ttflist:
    print(f.name)
"""


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


# Exception
def getArea(radius):
    if radius < 0:
        raise RuntimeError("Negative radius")

    return radius * radius * 3.1415


try:
    print(getArea(5))
    print(getArea(-5))
except RuntimeError:
    print("Invalid radius")

print("------------------------------------------------------------")  # 60個

# Create a deck of cards
deck = [x for x in range(0, 52)]

# Create suits and ranks lists
suits = ["Spades", "Hearts", "Diamonds", "Clubs"]
ranks = ["Ace", "2", "3", "4", "5", "6", "7", "8", "9", "10", "Jack", "Queen", "King"]

# Shuffle the cards
import random

random.shuffle(deck)

# Display the first four cards
for i in range(4):
    suit = suits[deck[i] // 13]
    rank = ranks[deck[i] % 13]
    print("Card number", deck[i], "is", rank, "of", suit)

print("------------------------------------------------------------")  # 60個

from random import randint

# Open file for writing data
outfile = open("tmp_Numbers.txt", "w")
for i in range(10):
    outfile.write(str(randint(0, 9)) + " ")
outfile.close()  # Close the file

# Open file for reading data
infile = open("tmp_Numbers.txt", "r")
s = infile.read()
numbers = [eval(x) for x in s.split()]
for number in numbers:
    print(number, end=" ")
infile.close()  # Close the file

print()

print("------------------------------------------------------------")  # 60個

import math  # import Math module to use the math functions

# Test algebraic functions
print("exp(1.0) =", math.exp(1))
print("log(3.78) =", math.log(math.e))
print("log10(10, 10) =", math.log(10, 10))
print("sqrt(4.0) =", math.sqrt(4.0))

# Test trigonometric functions
print("sin(PI / 2) =", math.sin(math.pi / 2))
print("cos(PI / 2) =", math.cos(math.pi / 2))
print("tan(PI / 2) =", math.tan(math.pi / 2))
print("degrees(1.57) =", math.degrees(1.57))
print("radians(90) =", math.radians(90))

print("------------------------------------------------------------")  # 60個
"""
#data analysis

NUMBER_OF_ELEMENTS = 5 # For simplicity, use 5 instead of 100
numbers = [] # Create an empty list
sum = 0

for i in range(NUMBER_OF_ELEMENTS): 
    value = eval(input("Enter a new number: "))
    numbers.append(value)
    sum += value
    
average = sum / NUMBER_OF_ELEMENTS

count = 0 # The number of elements above average
for i in range(NUMBER_OF_ELEMENTS): 
    if numbers[i] > average:
        count += 1

print("Average is", average)
print("Number of elements above the average is", count)
"""
print("------------------------------------------------------------")  # 60個

filename = "tmp_Presidents.txt"

# 製作一個檔案
# Open file for output
outfile = open(filename, "w")

# Write data to the file
outfile.write("Bill Clinton\n")
outfile.write("George Bush\n")
outfile.write("Barack Obama")

outfile.close()  # Close the output file

print("------------------------------------------------------------")  # 60個

filename = "tmp_Presidents.txt"

fp = open(filename, "r")
zops = fp.readlines()
fp.close()

i = 1
print("檔案內容")
for zen in zops:
    print("第 {} 行 : {}".format(i, zen), end="")
    i += 1

print()

print("------------------------------------------------------------")  # 60個

filename = "tmp_Presidents.txt"


def main():
    # Open file for input
    infile = open(filename, "r")
    print("(1) Using read(): ")
    print(infile.read())
    infile.close()  # Close the input file

    # Open file for input
    infile = open(filename, "r")
    print("\n(2) Using read(number): ")
    s1 = infile.read(4)
    print(s1)
    s2 = infile.read(10)
    print(repr(s2))
    infile.close()  # Close the input file

    # Open file for input
    infile = open(filename, "r")
    print("\n(3) Using readline(): ")
    line1 = infile.readline()
    line2 = infile.readline()
    line3 = infile.readline()
    line4 = infile.readline()
    print(repr(line1))
    print(repr(line2))
    print(repr(line3))
    print(repr(line4))
    infile.close()  # Close the input file

    # Open file for input
    infile = open(filename, "r")
    print("\n(4) Using readlines(): ")
    print(infile.readlines())
    infile.close()  # Close the input file


main()  # Call the main function

print("------------------------------------------------------------")  # 60個

print("各種讀取檔案的方法")

filename = "tmp_Presidents.txt"

# Open file for input
infile = open(filename, "r")
print("(1) Using read(): ")
print(infile.read())
infile.close()  # Close the input file

# Open file for input
infile = open(filename, "r")
print("\n(2) Using read(number): ")
s1 = infile.read(4)
print(s1)
s2 = infile.read(10)
print(repr(s2))
infile.close()  # Close the input file

# Open file for input
infile = open(filename, "r")
print("\n(3) Using readline(): ")
line1 = infile.readline()
line2 = infile.readline()
line3 = infile.readline()
line4 = infile.readline()
print(repr(line1))
print(repr(line2))
print(repr(line3))
print(repr(line4))
infile.close()  # Close the input file

# Open file for input
infile = open(filename, "r")
print("\n(4) Using readlines(): ")
print(infile.readlines())
infile.close()  # Close the input file

print("------------------------------------------------------------")  # 60個

import time

currentTime = time.time()  # Get current time

# Obtain the total seconds since midnight, Jan 1, 1970
totalSeconds = int(currentTime)

print(totalSeconds)

print("------------------------------------------------------------")  # 60個

""" 久
import urllib.request
input = urllib.request.urlopen('http://www.yahoo.com/index.html')
print(input.read())
"""

print("------------------------------------------------------------")  # 60個

# Social Security Number

import re

regex = "\d{3}-\d{2}-\d{4}"
# ssn = input("Enter SSN: ")
ssn = "123-45-6789"
match1 = re.match(regex, ssn)

if match1 != None:
    print(ssn, " is a valid SSN")
    print("start position of the matched text is " + str(match1.start()))
    print("start and end position of the matched text is " + str(match1.span()))
else:
    print(ssn, " is not a valid SSN")

print("------------------------------------------------------------")  # 60個

import re

regex = "\d{3}-\d{2}-\d{4}"
# text = input("Enter a text: ")
text = "123-45-6789"

match1 = re.search(regex, text)

if match1 != None:
    print(text, " contains a SSN")
    print("start position of the matched text is " + str(match1.start()))
    print("start and end position of the matched text is " + str(match1.span()))
else:
    print(text, " does not contain a SSN")

print("------------------------------------------------------------")  # 60個

s = "1 2 3 4 5"
items = s.split()  # Extract items from the string
lst = [eval(x) for x in items]  # Convert items to numbers

print(lst)

print("------------------------------------------------------------")  # 60個

students = [("John", "Smith", 96), ("Susan", "King", 76), ("Kim", "Yao", 99)]
students.sort(key=lambda e: (e[1]))

print(students)
print(sorted(students, key=lambda t: (t[2]), reverse=True))

print("------------------------------------------------------------")  # 60個

# 河內塔


# The function for finding the solution to move n disks
#   from fromTower to toTower with auxTower
def moveDisks(n, fromTower, toTower, auxTower):
    if n == 1:  # Stopping condition
        print("Move disk", n, "from", fromTower, "to", toTower)
    else:
        moveDisks(n - 1, fromTower, auxTower, toTower)
        print("Move disk", n, "from", fromTower, "to", toTower)
        moveDisks(n - 1, auxTower, toTower, fromTower)


print("盤子數 5")
n = 5

# Find the solution recursively
print("The moves are:")
moveDisks(n, "A", "B", "C")

print("------------------------------------------------------------")  # 60個


# Convert a decimal to a hex as a string
def decimalToHex(decimalValue):
    hex = ""

    while decimalValue != 0:
        hexValue = decimalValue % 16
        hex = toHexChar(hexValue) + hex
        decimalValue = decimalValue // 16

    return hex


# Convert an integer to a single hex digit in a character
def toHexChar(hexValue):
    if 0 <= hexValue <= 9:
        return chr(hexValue + ord("0"))
    else:  # 10 <= hexValue <= 15
        return chr(hexValue - 10 + ord("A"))


decimalValue = 170

print("The hex number for decimal", decimalValue, "is", decimalToHex(decimalValue))

print("------------------------------------------------------------")  # 60個


print("          九九乘法表")
# Display the number title
print("  |", end="")
for j in range(1, 10):
    print("  ", j, end="")
print()  # Jump to the new line
print("-----------------------------------------")

# Display table body
for i in range(1, 10):
    print(i, "|", end="")
    for j in range(1, 10):
        # Display the product and align properly
        print(format(i * j, "4d"), end="")
    print()  # Jump to the new line


print("------------------------------------------------------------")  # 60個


# Compute the distance between two points (x1, y1) and (x2, y2)
def distance(x1, y1, x2, y2):
    return ((x2 - x1) * (x2 - x1) + (y2 - y1) * (y2 - y1)) ** 0.5


def nearestPoints(points):
    # p1 and p2 are the indices in the points list
    p1, p2 = 0, 1  # Initial two points

    shortestDistance = distance(
        points[p1][0], points[p1][1], points[p2][0], points[p2][1]
    )  # Initialize shortestDistance

    # Compute distance for every two points
    for i in range(len(points)):
        for j in range(i + 1, len(points)):
            d = distance(
                points[i][0], points[i][1], points[j][0], points[j][1]
            )  # Find distance

            if shortestDistance > d:
                p1, p2 = i, j  # Update p1, p2
                shortestDistance = d  # New shortestDistance

    return p1, p2


print("找出多點中最近的兩點")

# Create a list to store points
points = []

point = 2 * [0]
point[0], point[1] = 0, 0
points.append(point)

point = 2 * [0]
point[0], point[1] = 5, 0
points.append(point)

point = 2 * [0]
point[0], point[1] = 5, 5
points.append(point)

point = 2 * [0]
point[0], point[1] = 3, 1
points.append(point)

# p1 and p2 are the indices in the points list
p1, p2 = nearestPoints(points)

# Display result
print(
    "The closest two points are ("
    + str(points[p1][0])
    + ", "
    + str(points[p1][1])
    + ") and ("
    + str(points[p2][0])
    + ", "
    + str(points[p2][1])
    + ")"
)

print("------------------------------------------------------------")  # 60個

import logging

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(message)s")

logging.info("Start to load dataset")

logging.info("Done with load dataset")

print("------------------------------------------------------------")  # 60個

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from io import BytesIO
from lxml import etree
import base64


df = pd.DataFrame(
    {
        "id": ["1", "2", "3", "4", "5"],  # 構造數據
        "math": [90, 89, 99, 78, 63],
        "english": [89, 94, 80, 81, 94],
    }
)


x = np.linspace(0, 8, 100)
y = np.sin(x)
plt.plot(x, y)

# df資料存成html
buffer = BytesIO()
plt.savefig(buffer)
plot_data = buffer.getvalue()

imb = base64.b64encode(plot_data)  # 生成網頁內容
ims = imb.decode()
imd = "data:image/png;base64," + ims
data_im = """<h1>Figure</h1>  """ + """<img src="%s">""" % imd
data_des = """<h1>Describe</h1>""" + df.describe().T.to_html()
root = "<title>Dataset</title>"
root = root + data_des + data_im

html = etree.HTML(root)
tree = etree.ElementTree(html)
tree.write("tmp_導出圖表.html")

plt.show()

print("------------------------------------------------------------")  # 60個

root_dir = os.path.abspath(".")
gunfire_path = os.path.join(root_dir, "gunfire.wav")
filename = os.path.join(root_dir, "tone.wav")

print(filename)


import os

"""
# 新建資料夾用於放置影像
if not os.path.isdir('trainer'):
    os.mkdir('trainer')
os.chdir('trainer')
"""

root_dir = os.path.abspath(".")
filename = os.path.join(root_dir, "tone.wav")
print(filename)

print("------------------------------------------------------------")  # 60個

print("建立20組資料")
USERS = [(i, f"account_{i}", f"nickname_{i}") for i in range(20)]

print(type(USERS))
print(len(USERS))
print(USERS)
for _ in USERS:
    print(_)

print("------------------------------------------------------------")  # 60個

"""
import glob,cv2

foldername = "animal"
 
#建立測試資料
filenames = glob.glob(foldername + '/*.jpg') + glob.glob(foldername + '/*.png')

for filename in filenames:
    print(filename)



print("------------------------------------------------------------")  # 60個

filename = 'C:/_git/vcs/_4.python/_data/picture1.jpg'

#filename = 'C:/_git/vcs/_4.python/_data/picture1.jpg'

os.system(filename)  # 用系統內建的程式開啟檔案

"""

print("------------------------------------------------------------")  # 60個

print("統計一串英文字串個字母出現的頻率")

from collections import defaultdict

text = "this is a lion-mouse"

frequency = defaultdict(int)
for symbol in text:
    frequency[symbol] += 1
print(frequency)

heap = [[weight, [symbol, ""]] for symbol, weight in frequency.items()]
print(heap)

print("------------------------------------------------------------")  # 60個

print("------------------------------------------------------------")  # 60個

"""
from pathlib import Path

# Version
BASE_DIR = Path(__file__).resolve().parent
print(BASE_DIR)

VERSION_FILE = BASE_DIR.joinpath('data_store', 'cccc.py')
print(VERSION_FILE)

#執行一個檔案
with open(VERSION_FILE) as fp:
    exec(fp.read())

import os
from pathlib import Path
from setuptools import setup, find_packages

import dicom

BASE_DIR = Path(__file__).parent
with open(BASE_DIR / "pydicom" / "_version.py") as f:
    exec(f.read())

with open(BASE_DIR / 'README.md') as f:
    long_description = f.read()

def data_files_inventory():
    root = BASE_DIR / "pydicom" / "data"
    files = [
        f.relative_to(BASE_DIR / "pydicom")
        for f in root.glob("**/*")
        if f.is_file() and f.suffix != ".pyc"
    ]
    return [os.fspath(f) for f in files]
"""

print("------------------------------------------------------------")  # 60個

from pathlib import Path

data_path = Path(__file__).resolve().parent.parent / "data"
print(data_path)

print("------------------------------------------------------------")  # 60個

import pydicom.data
from pydicom.data import get_testdata_file

fname = "693_UNCI.dcm"
fpath = get_testdata_file(fname)
print(fpath)

print("done")

print("------------------------------------------------------------")  # 60個

# UTC时间

import datetime

# 创建一个时间戳（以秒为单位）
timestamp = 22
# 带UTC时区时间
dt_with_timezone = datetime.datetime.fromtimestamp(timestamp, tz=datetime.timezone.utc)
print("带UTC时区时间:", dt_with_timezone)
# 不带UTC时区时间
dt_without_timezone = datetime.datetime.fromtimestamp(timestamp)
print("不带UTC时区时间", dt_without_timezone)

# 时间戳
import time

print(time.time())
print(time.localtime())  # 获取到当前时间的元组
print(time.mktime(time.localtime()))
# 一周的第几天(周一是0,0-6)、一年的第几天(从1开始，1-366)、夏时令(是夏时令1，不是0，未知-1)。

# 字符串和时间转换

# 字符串和时间转换
# 利用time模块的strftime()函数可以将时间戳转换成系统时间。
import time

time_str = time.strftime(("%Y-%m-%d %H:%M:%S"), time.localtime())
print(time_str)

# 可以用strptime函数将日期字符串转换为datetime数据类型，
import datetime

print(datetime.datetime.strptime("2022-01-15", "%Y-%m-%d"))

# 可以用Pandas的to_datetime()函数将日期字符串转换为datetime数据类型。
# to_datetime()函数转化后的时间是精准到时分秒精度的
import pandas as pd

print(pd.to_datetime("2022/01/15"))

# 时间差

# 3. 时间运算--时间差
# 利用datetime将时间类型数据进行转换，然后利用减法运算计算时间的不同之处
# 默认输出结果转换为用（“天”，“秒”）表达
import datetime

delta = datetime.datetime(2022, 1, 16) - datetime.datetime(2021, 1, 1, 9, 15)
print(delta)
print(delta.days)
print(delta.seconds)

print("------------------------------------------------------------")  # 60個

import numpy as np

x = np.array([1, 2, 3, 4, 1, 2, 3, 4, 5]).reshape(3, 3)
y = np.arange(9).reshape(3, 3)
print(x)
print(y)
print(x @ y)  # 不知道這是什麼運算

# 使用NumPy進行點積運算 ??

print("------------------------------------------------------------")  # 60個

import numpy as np

a = np.array([[1, 1], [10, 25]])
b = np.array([16, 250])
print(np.linalg.solve(a, b))

print("------------------------------------------------------------")  # 60個

print("數字對應 0~9 對應到 9~0")
old = np.array([3, 8, 3, 4, 2, 1, 4, 1, 2, 2])
print(old)

# 數字對應
new = np.choose(old, [9, 8, 7, 6, 5, 4, 3, 2, 1, 0]).astype(np.int64)

print(new)


print("------------------------------------------------------------")  # 60個

# 05_08_regularization
# L1、L2 regularization的計算與強度比較

# 權重
W = np.array([-1, 5, 3, -9]).reshape(2, 2)
print(W)

# L1

Lambda = 0.5
L1 = Lambda * np.sum(np.abs(W))
print(L1)

# L2

L2 = Lambda * np.sum(W**2)
print(L2)

# 58.0

print("結論：L2 強度較大")


print("------------------------------------------------------------")  # 60個


import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
import os

# 利用可能なカラーマップを取得
cmaps = plt.colormaps()
cmaps.sort()

# データの生成
xs = np.arange(1, 10)
ys = np.arange(1, 10).reshape(9, 1)
m = xs * ys
df = pd.DataFrame(m)


# テンプレートの読み込み
def load_template(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


# テンプレートを読み込む
header_template = load_template("templates/header_template.txt")
section_template = load_template("templates/section_template.txt")

# README.mdを作成
with open("README.md", "w", encoding="utf-8") as readme:
    readme.write(header_template)

    for cmap in cmaps:
        # ヒートマップを生成し、画像として保存
        plt.figure(figsize=(5, 3))
        ax = sns.heatmap(df, cmap=cmap)
        ax.tick_params(axis="both", which="both", length=0)
        ax.set_xticklabels([])
        ax.set_yticklabels([])
        ax.set_xlabel("")
        ax.set_ylabel("")
        plt.tight_layout(pad=0.1)
        plt.savefig(f"images/{cmap}.png", transparent=True)
        plt.close()

        # READMEにセクションを追加
        section_content = section_template.format(cmap_name=cmap)
        readme.write(section_content)

print("------------------------------------------------------------")  # 60個

print("列出所有區域變數的名稱與內容")
cc = locals()
print(cc)

print("列出所有全域變數的名稱與內容")
cc = globals()
print(cc)

import sys

# 目前 python程式 路徑
print(sys.executable)

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個
print("作業完成")
print("------------------------------------------------------------")  # 60個
sys.exit()

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


# Classes and Objects


class Dog:
    def __init__(self, name, age):
        self.name = name
        self.age = age

    def bark(self):
        return "Woof!"


# Create an instance of the Dog class
my_dog = Dog("Buddy", 3)

# Print the attributes of the instance
print("my_dog name:", my_dog.name)
print("my_dog age:", my_dog.age)

# Call the bark method of the instance
bark_result = my_dog.bark()
print("bark_result:", bark_result)

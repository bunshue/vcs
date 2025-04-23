import os
import sys

print("------------------------------------------------------------")  # 60個

print('搜尋字串')
keyword = 'shutil'
print("搜尋字串：{}".format(keyword))

cur_path = os.path.dirname(__file__) # 取得目前路徑
print(cur_path)
sample_tree = os.walk(cur_path)

for dirname, subdir, files in sample_tree:
   allfiles = []
   for file in files:  # 取得所有 .py .txt 檔，存入 allfiles 串列中
      print(file)
      ext = file.split('.')[-1]
      if ext == "py" or ext == "txt":
         allfiles.append(dirname + '/' + file)
         
   if len(allfiles) > 0:
      for file in allfiles:  # 讀取 allfiles 串列所有檔案
         try:
            fp = open(file, "r", encoding = 'UTF-8')
            article = fp.readlines()
            fp.close
            line = 0
            for row in article:
               line += 1
               if keyword in row:
                   print("在 {}，第 {} 列找到「{}」。".format(file, line, keyword))
         except:
            print("{} 無法讀取..." .format(file)) 
     
print("完成...")

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# FindKeyWord2.py

import docx

print('搜尋字串')
keyword = "籃球"
print("搜尋字串：{}".format(keyword))

cur_path = os.path.dirname(__file__) # 取得目前路徑
print(cur_path)
sample_tree = os.walk(cur_path)

for dirname, subdir, files in sample_tree:
   allfiles = []   
   for file in files:  # 取得所有 .docx 檔，存入 allfiles 串列中
      print(file)
      ext = file.split('.')[-1]
      if ext == "docx": # get *.docx to allfiles
         allfiles.append(dirname + '/' + file)
         
   for file in allfiles:
      print("正在搜尋 <{}> 檔案...".format(file))
      try:
         doc = docx.Document(file)
         line = 0
         for p in doc.paragraphs:
             line += 1
             if keyword in p.text:
                 print("...在第 {} 段文字中找到「{}」\n {}。".format(line, keyword, p.text))
      except:
         print("無法讀取 {} 檔案..." .format(file))  
     
print("\n搜尋完畢...")

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# FindKeyWord3.py

import docx

print('搜尋字串')
keyword = 'shutil'
print("搜尋字串：{}".format(keyword))

#cur_path = os.path.dirname(__file__) # 取得目前路徑
cur_path = os.getcwd()
print(cur_path)
sample_tree = os.walk(cur_path)

for dirname, subdir, files in sample_tree:
   allfiles = []   
   for file in files:  # 取得所有 .py .txt .docx 檔，存入 allfiles 串列中
      print(file)
      ext = file.split('.')[-1]
      if ext == "py" or ext == "txt" or ext == "docx": 
         allfiles.append(dirname + '/' + file)
         
   if len(allfiles) > 0:  
      for file in allfiles:  # 讀取 allfiles 串列所有檔案  
         try:
            if file.split('.')[-1] == "docx": # .docx
                doc = docx.Document(file)
                line = 0
                for p in doc.paragraphs:
                    line += 1
                    if keyword in p.text:
                        print("...在第 {} 段文字中找到「{}」\n {}。".format(line, keyword, p.text))
            else:  # .py or .txt             
                fp = open(file, "r", encoding = 'UTF-8')
                article = fp.readlines()
                fp.close 
                line = 0            
                for row in article:
                   line += 1
                   if keyword in row:
                       print("在 {}，第 {} 列找到「{}」。".format(file, line, keyword))
         except:
            print("{} 無法讀取..." .format(file)) 
     
print("完成...")

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

"""
模式有
r - 讀取(檔案需存在)
w - 新建檔案寫入(檔案可不存在，若存在則清空)
a - 資料附加到舊檔案後面(游標指在EOF)
"""

#附加模式
f = open('tmp_A.txt', 'a', encoding = 'UTF-8')   # 也可使用指定路徑等方式，如： C:\A.txt
f.write('你好1\n')
f.write('你好2\n')
f.write('你好3\n')
f.close()


"""
Hello world
Today id a nice day!
"""


f=open('test_file.txt')
print(f.readline())
print(f.readline())
#f.close()


f.seek(0)
for line in f:
    print(line.strip())
    
f.close()

with open("test_file.txt") as f:
    for line in f:
        print(line.strip())          

# write to test_write.txt
# Write a file
with open("tmp_test_write.txt", "w") as out_file:
    out_file.write("This Text is going to out file\nLook at it and see!")

print("------------------------------------------------------------")  # 60個


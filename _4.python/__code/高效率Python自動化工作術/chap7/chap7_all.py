from pathlib import Path
from docx import Document
import re
import unicodedata

print("------------------------------------------------------------")  # 60個
print("find_Words")
print("------------------------------------------------------------")  # 60個

infolder = "testfolder"
value1 = "這個是"
value2 = "*.docx"


# 【函數: 搜尋Word檔案的字串】
def findfile(readfile, findword):
    try:
        msg = ""
        doc = Document(readfile)
        cnt = 0
        for pa in doc.paragraphs:  # 所有段落
            cnt += pa.text.count(findword)
        for tbl in doc.tables:  # 所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cnt += cell.text.count(findword)
        if cnt > 0:  # 找到的話
            msg = readfile + "：" + "找到" + str(cnt) + "個了。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"


# 【函數: 搜尋資料夾與子資料夾的文字檔】
def findfiles(infolder, findword, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).rglob(ext):  # 將這個資料夾以及子資料夾的所有檔案
        filelist.append(str(p))  # 新增至列表
    for filename in sorted(filelist):  # 再替每個檔案排序
        msg += findfile(filename, findword)
    return msg


# 【執行函數】
msg = findfiles(infolder, value1, value2)
print(msg)

print("------------------------------------------------------------")  # 60個
print("test7_1")
print("------------------------------------------------------------")  # 60個

infile = "test.docx"
try:
    doc = Document(infile)
    for pa in doc.paragraphs:  # 所有段落
        print("paragraph----")
        print(pa.text)
    for tbl in doc.tables:  # 所有表
        print("table----")
        for row in tbl.rows:
            print("row----")
            for cell in row.cells:
                print(cell.text)
except:
    print("程式執行失敗。")


print("------------------------------------------------------------")  # 60個

infile = "test.docx"
value1 = "這個是"

try:
    doc = Document(infile)
    cnt = 0
    for pa in doc.paragraphs:  # 所有段落
        cnt += pa.text.count(value1)
    for tbl in doc.tables:  # 所有表
        for row in tbl.rows:
            for cell in row.cells:
                cnt += cell.text.count(value1)
    print("找到" + str(cnt) + "個了。")
except:
    print("程式執行失敗。")


print("------------------------------------------------------------")  # 60個

infile = "test.docx"
value1 = "tmp_output.docx"
try:
    doc = Document(infile)
    doc.save(value1)
except:
    print("程式執行失敗。")

print("------------------------------------------------------------")  # 60個
print("regfind_Words")
print("------------------------------------------------------------")  # 60個

infolder = "testfolder"
value1 = ".個是"
value2 = "*.docx"


# 【函數: 搜尋Word檔案的字串】
def findfile(readfile, findword):
    try:
        msg = ""
        ptn = re.compile(findword)
        doc = Document(readfile)
        cnt = 0
        for pa in doc.paragraphs:  # 所有段落
            cnt += len(re.findall(ptn, pa.text))
        for tbl in doc.tables:  # 所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cnt += len(re.findall(ptn, cell.text))
        if cnt > 0:  # 找到的話
            msg = readfile + "：" + "找到" + str(cnt) + "個了。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"


# 【函數: 搜尋資料夾與子資料夾的文字檔】
def findfiles(infolder, findword, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).rglob(ext):  # 將這個資料夾以及子資料夾的所有檔案
        filelist.append(str(p))  # 新增至列表
    for filename in sorted(filelist):  # 再替每個檔案排序
        msg += findfile(filename, findword)
    return msg


# 【執行函數】
msg = findfiles(infolder, value1, value2)
print(msg)

print("------------------------------------------------------------")  # 60個
print("replace_Words")
print("------------------------------------------------------------")  # 60個

infolder = "testfolder"
value1 = "這個是"
value2 = "那個是"
value3 = "outputfolder"
ext = "*.docx"


# 【函數: 置換文字檔】
def replacefile(readfile, findword, newword, savefolder):
    try:
        msg = ""
        doc = Document(readfile)
        for pa in doc.paragraphs:  # 所有段落
            pa.text = pa.text.replace(findword, newword)
        for tbl in doc.tables:  # 所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace(findword, newword)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)  # 建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)  # Word轉存檔案
        msg = "在" + savefolder + "轉存" + filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"


# 【函數：置換資料夾之內的文字檔】
def replacefiles(infolder, findword, newword, savefolder):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  # 將這個資料夾的檔案
        filelist.append(str(p))  # 新增至列表
    for filename in sorted(filelist):  # 再替每個檔案排序
        msg += replacefile(filename, findword, newword, savefolder)
    return msg


# 【執行函數】
msg = replacefiles(infolder, value1, value2, value3)
print(msg)


print("------------------------------------------------------------")  # 60個
print("regreplace_Words")
print("------------------------------------------------------------")  # 60個

infolder = "testfolder"
value1 = ".個是"
value2 = "那個是"
value3 = "outputfolder"
ext = "*.docx"


# 【函數: 置換文字檔】
def replacefile(readfile, findword, newword, savefolder):
    try:
        msg = ""
        ptn = re.compile(findword)
        doc = Document(readfile)
        for pa in doc.paragraphs:  # 所有段落
            pa.text = re.sub(ptn, newword, pa.text)
        for tbl in doc.tables:  # 所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = re.sub(ptn, newword, cell.text)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)  # 建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)  # Word轉存檔案
        msg = "在" + savefolder + "轉存" + filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"


# 【函數：置換資料夾之內的文字檔】
def replacefiles(infolder, findword, newword, savefolder):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  # 將這個資料夾的檔案
        filelist.append(str(p))  # 新增至列表
    for filename in sorted(filelist):  # 再替每個檔案排序
        msg += replacefile(filename, findword, newword, savefolder)
    return msg


# 【執行函數】
msg = replacefiles(infolder, value1, value2, value3)
print(msg)

print("------------------------------------------------------------")  # 60個
print("normalize_Words")
print("------------------------------------------------------------")  # 60個

infolder = "testfolder"
value1 = "outputfolder"
value2 = "*.docx"


# 【函數: 對Word檔案執行unicode正規化處理】
def normalizefile(readfile, savefolder):
    try:
        msg = ""
        doc = Document(readfile)
        for pa in doc.paragraphs:  # 所有段落
            pa.text = unicodedata.normalize("NFKC", pa.text)
        for tbl in doc.tables:  # 所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = unicodedata.normalize("NFKC", cell.text)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)  # 建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)  # Word轉存檔案
        msg = "在" + savefolder + "轉存" + filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"


# 【函數: 對資料夾之內的文字檔執行unicode正規化】
def normalizefiles(infolder, savefolder, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  # 將這個資料夾的檔案
        filelist.append(str(p))  # 新增至列表
    for filename in sorted(filelist):  # 再替每個檔案排序
        msg += normalizefile(filename, savefolder)
    return msg


# 【執行函數】
msg = normalizefiles(infolder, value1, value2)
print(msg)



print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個
print("作業完成")
print("------------------------------------------------------------")  # 60個
sys.exit()

print("------------------------------------------------------------")  # 60個


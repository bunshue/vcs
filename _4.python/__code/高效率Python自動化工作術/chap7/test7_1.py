from docx import Document

infile = "test.docx"
try:
    doc = Document(infile)
    for pa in doc.paragraphs:  # 所有段落
        print("paragraph----")
        print(pa.text)
    for tbl in doc.tables:  # 所有表
        print("table----")
        for row in tbl.rows:
            print("row----")
            for cell in row.cells:
                print(cell.text)
except:
    print("程式執行失敗。")


print("------------------------------------------------------------")  # 60個

from docx import Document

infile = "test.docx"
value1 = "這個是"

try:
    doc = Document(infile)
    cnt = 0
    for pa in doc.paragraphs:  # 所有段落
        cnt += pa.text.count(value1)
    for tbl in doc.tables:  # 所有表
        for row in tbl.rows:
            for cell in row.cells:
                cnt += cell.text.count(value1)
    print("找到" + str(cnt) + "個了。")
except:
    print("程式執行失敗。")


print("------------------------------------------------------------")  # 60個

from docx import Document

infile = "test.docx"
value1 = "tmp_output.docx"
try:
    doc = Document(infile)
    doc.save(value1)
except:
    print("程式執行失敗。")


print("------------------------------------------------------------")  # 60個

print("------------------------------------------------------------")  # 60個

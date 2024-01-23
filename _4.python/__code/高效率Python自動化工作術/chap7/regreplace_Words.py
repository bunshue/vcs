from pathlib import Path
from docx import Document
import re

infolder = "testfolder"
value1 = ".個是"
value2 = "那個是"
value3 = "outputfolder"
ext = "*.docx"

#【函數: 置換文字檔】
def replacefile(readfile, findword, newword, savefolder):
    try:
        msg = ""
        ptn = re.compile(findword)
        doc = Document(readfile)
        for pa in doc.paragraphs:               #所有段落
            pa.text = re.sub(ptn, newword, pa.text)
        for tbl in doc.tables:                  #所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = re.sub(ptn, newword, cell.text)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)            #建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)                       #Word轉存檔案
        msg = "在" + savefolder+"轉存"+ filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"
#【函數：置換資料夾之內的文字檔】
def replacefiles(infolder, findword, newword, savefolder):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  #將這個資料夾的檔案
        filelist.append(str(p))         #新增至列表
    for filename in sorted(filelist):   #再替每個檔案排序
        msg += replacefile(filename, findword, newword, savefolder)
    return msg

#【執行函數】
msg = replacefiles(infolder, value1, value2, value3)
print(msg)

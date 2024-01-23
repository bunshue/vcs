from pathlib import Path
from docx import Document
import re

infolder = "testfolder"
value1 = ".個是"
value2 = "*.docx"

#【函數: 搜尋Word檔案的字串】
def findfile(readfile, findword):
    try:
        msg = ""
        ptn = re.compile(findword)
        doc = Document(readfile)
        cnt = 0
        for pa in doc.paragraphs:           #所有段落
            cnt += len(re.findall(ptn, pa.text))
        for tbl in doc.tables:              #所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cnt += len(re.findall(ptn, cell.text))
        if cnt > 0:                         #找到的話
            msg = readfile+"："+"找到" + str(cnt)+"個了。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"
#【函數: 搜尋資料夾與子資料夾的文字檔】
def findfiles(infolder, findword, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).rglob(ext): #將這個資料夾以及子資料夾的所有檔案
        filelist.append(str(p))         #新增至列表
    for filename in sorted(filelist):   #再替每個檔案排序
        msg += findfile(filename, findword)
    return msg

#【執行函數】
msg = findfiles(infolder, value1, value2)
print(msg)

import PySimpleGUI as sg
#--------------------vvv
#【1. import函式庫】
from pathlib import Path
import unicodedata
from docx import Document

#【2. 設定於應用程式顯示的字串】
title = "對Word檔案進行unicode正規化處理（資料夾之內的文字檔）"
infolder = "."
label1, value1 = "轉存資料夾", "outputfolder"
label2, value2 = "副檔名", "*.docx"

#【3.函數：執行unicode正規化】
def normalizefile(readfile, savefolder):
    try:
        msg = ""
        doc = Document(readfile)
        for pa in doc.paragraphs:               #所有段落
            pa.text = unicodedata.normalize("NFKC", pa.text)
        for tbl in doc.tables:                  #所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = unicodedata.normalize("NFKC", cell.text)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)            #建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)                       #Word轉存檔案
        msg = "在" + savefolder+"轉存"+ filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"
#【3.函數: 對資料夾之內的文字檔執行unicode正規化】
def normalizefiles(infolder, savefolder, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  #將這個資料夾的檔案
        filelist.append(str(p))         #新增至列表
    for filename in sorted(filelist):   #再替每個檔案排序
        msg += normalizefile(filename, savefolder)
    return msg
#--------------------^^^
def execute():
    infolder = values["infolder"]
    value1 = values["input1"]
    #--------------------vvv
    #【4.執行函數】
    msg = normalizefiles(infolder, value1, value2)
    #--------------------^^^
    window["text1"].update(msg)
#應用程式的介面
layout = [[sg.Text("要載入的資料夾", size=(14,1)),
           sg.Input(infolder, key="infolder"),sg.FolderBrowse("選取")],
          [sg.Text(label1, size=(14,1)), sg.Input(value1, key="input1")],
          [sg.Button("執行", size=(20,1), pad=(5,15), bind_return_key=True)],
          [sg.Multiline(key="text1", size=(60,10))]]
#執行應用程式的處理
window = sg.Window(title, layout, font=(None,14))
while True:
    event, values = window.read()
    if event == None:
        break
    if event == "執行":
      execute()
window.close()

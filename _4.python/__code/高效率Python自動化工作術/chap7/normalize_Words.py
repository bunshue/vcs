from pathlib import Path
import unicodedata
from docx import Document

infolder = "testfolder"
value1 = "outputfolder"
value2 = "*.docx"

#【函數: 對Word檔案執行unicode正規化處理】
def normalizefile(readfile, savefolder):
    try:
        msg = ""
        doc = Document(readfile)
        for pa in doc.paragraphs:               #所有段落
            pa.text = unicodedata.normalize("NFKC", pa.text)
        for tbl in doc.tables:                  #所有表
            for row in tbl.rows:
                for cell in row.cells:
                    cell.text = unicodedata.normalize("NFKC", cell.text)
        savedir = Path(savefolder)
        savedir.mkdir(exist_ok=True)            #建立轉存資料夾
        filename = Path(readfile).name
        newname = savedir.joinpath(filename)
        doc.save(newname)                       #Word轉存檔案
        msg = "在" + savefolder+"轉存"+ filename + "了喲。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"
#【函數: 對資料夾之內的文字檔執行unicode正規化】
def normalizefiles(infolder, savefolder, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).glob(ext):  #將這個資料夾的檔案
        filelist.append(str(p))         #新增至列表
    for filename in sorted(filelist):   #再替每個檔案排序
        msg += normalizefile(filename, savefolder)
    return msg

#【執行函數】
msg = normalizefiles(infolder, value1, value2)
print(msg)
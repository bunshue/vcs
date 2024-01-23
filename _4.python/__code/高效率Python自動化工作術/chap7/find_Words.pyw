import PySimpleGUI as sg
#--------------------vvv
#【1. import函式庫】
from pathlib import Path
from docx import Document

#【2. 設定於應用程式顯示的字串】
title = "搜尋Word檔（資料夾與子資料夾）"
infolder = "."
label1, value1 = "要搜尋的字串", "這個是"
label2, value2 = "副檔名", "*.docx"

#【3.函數：搜尋文字檔】
def findfile(readfile, findword):
    try:
        msg = ""
        doc = Document(readfile)
        cnt = 0
        for pa in doc.paragraphs:           #所有段落
          cnt += pa.text.count(findword)
        for tbl in doc.tables:              #所有表
          for row in tbl.rows:
              for cell in row.cells:
                cnt += cell.text.count(findword)
        if cnt > 0:                         #找到的話
            msg = readfile+"："+"找到" + str(cnt)+"個了。\n"
        return msg
    except:
        return readfile + "：程式執行失敗。"
#【3.函數: 搜尋資料夾與子資料夾的文字檔】
def findfiles(infolder, findword, ext):
    msg = ""
    filelist = []
    for p in Path(infolder).rglob(ext): #將這個資料夾以及子資料夾的所有檔案
        filelist.append(str(p))         #新增至列表
    for filename in sorted(filelist):   #再替每個檔案排序
        msg += findfile(filename, findword)
    return msg
#--------------------^^^
def execute():
    infolder = values["infolder"]
    value1 = values["input1"]
    #--------------------vvv
    #【4.執行函數】
    msg = findfiles(infolder, value1, value2)
    #--------------------^^^
    window["text1"].update(msg)
#應用程式的介面
layout = [[sg.Text("要載入的資料夾", size=(14,1)),
           sg.Input(infolder, key="infolder"),sg.FolderBrowse("選取")],
          [sg.Text(label1, size=(14,1)), sg.Input(value1, key="input1")],
          [sg.Button("執行", size=(20,1), pad=(5,15), bind_return_key=True)],
          [sg.Multiline(key="text1", size=(60,10))]]
#執行應用程式的處理
window = sg.Window(title, layout, font=(None,14))
while True:
    event, values = window.read()
    if event == None:
        break
    if event == "執行":
      execute()
window.close()

#各種檔案寫讀範例 新進暫存

print('------------------------------------------------------------')	#60個

import sys, ast

filename = 'C:/_git/vcs/_1.data/______test_files1/__RW/_csv/scores.csv'

scores = dict()
with open(filename,'r') as fp:
    filedata = fp.read()
    #scores = ast.literal_eval(filedata)
print("以下是{}成績檔的字典型態資料:".format(filename))

print('------------------------------------------------------------')	#60個

"""
import sys

std_data = dict()
with open(filename, encoding='utf-8') as fp:
    alldata = fp.readlines()
    for item in alldata:
        no, name = item.rstrip('\n').split(',')
        std_data[no] = name
print(std_data)
"""

print('------------------------------------------------------------')	#60個

"""
#readlines()可以依照行讀取整個檔案，回傳是一個List，每一個element就是一行字。

file = open("demo_en.txt", "r")

lines = file.readlines()
print(lines)

file.close()
"""

print('------------------------------------------------------------')	#60個

"""
print('一種寫入檔案的方法')
filename = 'tmp.txt'

fp = open(filename,'w')
print('[OPTIONS]', file=fp)
print('Auto Index=Yes', file=fp)
print('Binary TOC=No', file=fp)
print('Binary Index=Yes', file=fp)
print('Compatibility=1.1', file=fp)
print('Error log file=ErrorLog.log', file=fp)
print('Display compile progress=Yes', file=fp)
print('Full-text search=Yes', file=fp)
print('Default window=main', file=fp)
print('', file=fp)  #寫一個空白列
print('[WINDOWS]', file=fp)
print('main=,"' + '","'
+ '","","",,,,,0x23520,222,0x1046,[10,10,780,560],'
'0xB0000,,,,,,0', file=fp)
print('', file=fp)
print('[FILES]', file=fp)
print('', file=fp)
fp.close()
"""



"""
json

filename = 'C:/_git/vcs/_1.data/______test_files2/news_' + time.strftime("%Y%m%d_%H%M%S", time.localtime()) + '.json';
with open(filename, "w", encoding = 'utf-8') as fp:
    print(filename + " is dumping...")
    json.dump(titles, fp)
"""



"""
print('11')
filename = 'engnews.txt'
with open(filename, "r", encoding="utf-8") as f:
    print(f.readline())#讀一行

print('22')
filename = 'engnews.txt'    
with open(filename, "r", encoding="utf-8") as f:
    data = f.read() #讀全部成一行串列

print(repr(data))
print(data)
print(data.split())
data = data.split()
for d in data:
    d.strip()
print(data)

print('33')
filename = 'engnews.txt'    
with open(filename, "r", encoding="utf-8") as f:
    print(f.readlines())#讀全部成多行串列
"""

print("------------------------------------------------------------")  # 60個


"""
file_Obj =  open(fn, encoding='cp950')  # 用預設encoding='cp950'開啟檔案
file_Obj =  open(fn, encoding='cp950')  # 用預設encoding='cp950'開啟檔案
file_Obj =  open(fn, encoding='utf-8')  # 用encoding='utf-8'開啟檔案
with open(fn, encoding='utf-8') as file_Obj:    # 開啟utf-8檔案
    obj_list = file_Obj.readlines()             # 每次讀一行
with open(fn, encoding='utf-8-sig') as file_Obj:  # 開啟utf-8檔案
    obj_list = file_Obj.readlines()               # 每次讀一行
"""

print("------------------------------------------------------------")  # 60個

file = open("data/temp.txt", "w+")

file.write("ABCDEFGHIJKLMNOPQRSTUVWXYZ")

file.flush()

print("寫入之後的游標位置：", file.tell())

print('往後跳16拜')
file.seek(16, 0)

print('擷取至位置26')
file.truncate(26)

print('讀出來')
print(file.read())

file1=open("data/temp.txt","r")
text=file1.read(1) #以read()方法讀取檔案
print(text)
text=file1.read(3) #以read()方法讀取檔案
print(text)
text=file1.read(2) #以read()方法讀取檔案
print(text)
text=file1.read(2) #以read()方法讀取檔案
print(text)
file1.close()


print("------------------------------------------------------------")  # 60個

print("讀出檔案最後一行字")

def read_final_line(filename):
    f = open(filename, 'r')
    for line in f:
        pass
    f.close()
    return line

print(read_final_line(r'.\data\login.log'))


print("------------------------------------------------------------")  # 60個


print('------------------------------------------------------------')	#60個

#練習 19 擷取登入帳號資訊

def passwd_to_dict(filename):
    users = {}
    with open(filename) as f:
        for line in f:
            user_info = line.split(':')
            users.update({user_info[0]: user_info[2]})
    return users

print(passwd_to_dict(r'.\data\passwd.cfg'))



print("------------------------------------------------------------")  # 60個


#練習 20 統計檔案的字元數、字數與行數

def wordcount(filename):
    result = {
        'Characters': 0,
        'Words': 0,
        'Unique words': 0,
        'Lines': 0,
        }
    unique_words = set()
 
    with open(filename, 'r') as f:
        for line in f:
            words = line.split()
            result['Lines'] += 1
            result['Characters'] += len(line)
            result['Words'] += len(words)
            unique_words.update(words)
 
        result['Unique words'] = len(unique_words)
 
    for key, value in result.items():
        print(f'{key}: {value}')

wordcount(r'.\data\text.txt')

print('------------------------------------------------------------')	#60個

#練習 21 找出檔案內的最長單字

def find_longest_word(filename):
    longest = ''
    with open(filename, 'r') as f:
        for line in f:
            for word in line.replace('.', '').split():
                if len(word) > len(longest):
                    longest = word
    return longest

print(find_longest_word(r'.\data\text2.txt'))

print('------------------------------------------------------------')	#60個

#練習 22 讀寫 CSV 檔

import csv

def passwd_to_csv(passwd_filename, csv_filename):
    with open(passwd_filename, 'r') as f_read, \
            open(csv_filename, 'w', newline='') as f_write:
        csv_reader = csv.reader(f_read, delimiter=':')
        csv_writer = csv.writer(f_write, delimiter='\t', lineterminator='\n')
        for line in csv_reader:
            csv_writer.writerow([line[0], line[2]])

passwd_to_csv(r'.\data\passwd.cfg', r'tmp_passwd.csv')

print('------------------------------------------------------------')	#60個


#練習 31 豬拉丁文 --- 檔案翻譯機

def pl_word(word):
    if word[0] in 'aeiou':
        return f'{word}way'
    return f'{word[1:]}{word[0]}ay'

def pl_file(filename):
    with open(filename, 'r') as f:
        return ' '.join([pl_word(word.lower().replace('.', ''))
                         for line in f
                         for word in line.split()])

print(pl_file(r'.\data\text2.txt'))

print('------------------------------------------------------------')	#60個


#練習 33 擷取登入帳號資訊（生成版）

def passwd_to_dict_2(filename):
    with open(filename) as f:
        d = {words[0]: words[2]
             for words
             in [line.split(':') for line in f]}
    return d

print(passwd_to_dict_2(r'.\data\passwd.cfg'))

print('------------------------------------------------------------')	#60個

#練習 34 過濾檔案中特定條件的單字

def word_filter(filename):
    vowels = {'a', 'e', 'i', 'o', 'u'}
    with open(filename, 'r') as f:
        words = ([word.replace('.', '')
                  for line in f
                      for word in line.split()
                      if len(set(word) & vowels) >= 3])
    return words

print(word_filter(r'.\data\text2.txt'))

print('------------------------------------------------------------')	#60個

#練習 35 希伯來數字密碼（Part I + Part II）

import string

def gematria_dict():
    return {char: index
            for index, char
            in enumerate(string.ascii_lowercase, 1)}

GEMATRIA = gematria_dict()

def gematria_value(word):
    return sum(GEMATRIA[char]
               for char in word.lower()
               if char in GEMATRIA)

def gematria_equal_words(my_word, filename):
    my_value = gematria_value(my_word)
    with open(filename, 'r', encoding='utf-8') as f:
        return [word
                for line in f
                for word in line.lower().split()
                if my_value == gematria_value(word)]

print(gematria_equal_words('programming', r'.\data\book.txt'))



print('------------------------------------------------------------')	#60個

#練習 48 檔案單字產生器

def word_generator(filename, max_words):
    index = 0
    with open(filename, 'r') as file:
        for line in file:
            for word in line.split():
                if index >= max_words:
                    return
                yield word.replace('.', '')
                index += 1

ten_words = word_generator(r'.\data\text2.txt', 10)

for word in ten_words:
    print(word)

print('------------------------------------------------------------')	#60個

print('直接 print 到檔案')

with open('tmp_cccccc.txt', 'wt') as f:
    print('Hello World!', file = f)

print('------------------------------------------------------------')	#60個


print('------------------------------------------------------------')	#60個


print("------------------------------------------------------------")  # 60個

filename = 'C:/_git/vcs/_4.python/_data/Romeo&Juliet.txt'

infile = open(filename, 'r')

count = 0
for line in infile:
    str1 = line.strip()
    if len(str1)>0:
        count += 1

print('total', count, 'lines')
infile.close()

print("------------------------------------------------------------")  # 60個

def distinctNN(n):
    set1 = set()
    for i in range(2, n+1):
        for j in range(1, n+1):
            set1.add(i*j)
    return set1

def doOutput(fname, outlist):
    outfile = open(fname, 'w')
    
    count = 0
    for num in outlist:
        outfile.write(str(num)+' ')
        count += 1
        if count%5==0:
            outfile.write('\n')
            
    outfile.close()
    print(fname+' created')

setNN = distinctNN(19)
list1 = sorted(list(setNN))
#print(list1)
outfname = 'tmp_SortedNN.txt'
doOutput(outfname, list1)

print("------------------------------------------------------------")  # 60個

"""
filename = 'myfile.txt'
inf = open(filename, 'r')
outfname = filename[:-4]+'2.txt'
outf = open(outfname, 'w')

for line in inf:           # 讀進來的line字串是有包含檔案內的換行字元哦！
    str1 = line.strip()    # 移除line的多餘空白
    if len(str1)>0:        # 如果移除完還有內容，寫進輸出檔
        outf.write(str1+'\n')
inf.close()
outf.close()
"""

print("------------------------------------------------------------")  # 60個

filename = 'data/en-us2.log'
infile = open(filename, 'r')

# 前50行
count = 0
for line in infile:
    if count>=50:
        break
    print((count+1), line, end='')
    count += 1
    
infile.close()


print("------------------------------------------------------------")  # 60個

""" many

filename = 'C:/_git/vcs/_4.python/_data/Romeo&Juliet.txt'
infile = open(filename, 'r')

count = 0
str1 = infile.readline()
len1 = len(str1)
while len1>0:
    count += 1
    print(count, str1.strip())
    str1 = infile.readline()
    len1 = len(str1)
    

print('total', count, 'lines')
infile.close()
"""

print("------------------------------------------------------------")  # 60個

filename = 'C:/_git/vcs/_1.data/______test_files1/picture1.jpg'

""" many
inf = open(filename, "rb")    # jpeg圖檔就是二進位檔案

byte = inf.read(1)               # 一次讀一個byte
while byte:
    print(byte)
    byte = inf.read(1)           # 繼續讀下一個byte

inf.close()
"""

print("------------------------------------------------------------")  # 60個


# BMP parser
def b4_2_int(bytes1):
    return (bytes1[0] | bytes1[1]<<8 | bytes1[2]<<16 | bytes1[3]<<24)

def b2_2_int(bytes1):
    return (bytes1[0] | bytes1[1]<<8)

filename = 'data/Medrust3.bmp'
infile = open(filename, 'rb')

# 'B' 'M'
b2 = infile.read(2)
print(b2)
b4 = infile.read(4)
size = b4_2_int(b4)
print('file size =', size, 'bytes')
infile.read(4)
infile.read(4)
infile.read(4)
width = b4_2_int(infile.read(4))
height = b4_2_int(infile.read(4))
print('image dimension =', width, 'x', height)
planes = b2_2_int(infile.read(2))
print('image planes =', planes)
bitsPerPixel = b2_2_int(infile.read(2))
print('bits per pixel =', bitsPerPixel)

infile.close()


print("------------------------------------------------------------")  # 60個

bytestr1 = b'This is a byte string'
int1 = 65
float1 = 3.14
int2s = [1, 2, 3]

outfname = 'tmp_test.bin'
outfile = open(outfname, 'wb')

outfile.write(bytestr1)
bint1 = int1.to_bytes(4, byteorder='big')
outfile.write(bint1)
bfloat1 = bytes(float1.hex(), 'utf-8')
outfile.write(bfloat1)
outfile.write(bytearray(int2s))

outfile.close()

print("------------------------------------------------------------")  # 60個

"""
with open("myfile.txt", 'r') as f:
    lines = f.readlines()
    for line in lines:
        print(line, end='')
        
# 不用關檔了！！
"""
print("------------------------------------------------------------")  # 60個


data = b'wxy\x7a'
print(data)               # b'wxyz'，以ASCII字元輸出

print(type(data), type(data[0]))
# <class 'bytes'>, <class 'int'>

print(data[0], hex(data[0]))
# 'w' ASCII碼119，十六進位'0x77'

print(b'\x7a' in data)    # 可以用 in 來判斷
print(data[2:])           # 可以切片

print("------------------------------------------------------------")  # 60個

data = b'wxy\x7a'
print(data)               # b'wxyz'，以ASCII字元輸出

ba = bytearray(data)
print(type(ba), type(ba[0]))
# <class 'bytearray'>, <class 'int'>

ba[3] = 0x70              # 修改資料
print(ba)                 # 變成 bytearray(b'wxyp')


print("------------------------------------------------------------")  # 60個

str1 = '葉'

print(str1.encode('big5'))     
# 轉big-5編碼，b'\xb8\xad'

print(str1.encode('gbk'))      
# 轉bgk編碼，b'\xc8~'

print(str1.encode('utf-8'))    
# 轉utf-8編碼，b'\xe8\x91\x89'


data = b'\xB8\xAD'

print(data.decode('big5'))
# 將bytes轉為big5文字，'葉'

print(data.decode('gbk'))
# 將bytes轉為gbk文字，'腑'

"""
print(data.decode('utf-8'))
# 將bytes轉為ytf-8文字，解碼規則無法解碼
"""

print("------------------------------------------------------------")  # 60個



import csv

filename = 'C:/_git/vcs/_1.data/______test_files1/__RW/_csv/animals.csv'

try:
    with open(filename, encoding = 'utf-8') as f:
        reader = csv.reader(f)
        data = list(reader)
except FileNotFoundError:
    print('无法打开文件:', filename)
else:
    for item in data:
        print('%-30s%-20s%-10s' % (item[0], item[1], item[2]))


print("------------------------------------------------------------")  # 60個


"""
从文本文件中读取数据

"""

""" fail
import time

def main():
    # 一次性读取整个文件内容
    with open('data/致橡树.txt', 'r', encoding='utf-8') as f:
        print(f.read())

    # 通过for-in循环逐行读取
    with open('data/致橡树.txt', mode='r') as f:
        for line in f:
            print(line, end='')
            time.sleep(0.5)
    print()

    # 读取文件按行读取到列表中
    with open('data/致橡树.txt') as f:
        lines = f.readlines()
    print(lines)

if __name__ == '__main__':
    main()
"""

print("------------------------------------------------------------")  # 60個

#檔案 : C:\_git\vcs\_4.python\__code\Python-100-Days-zh_TW-master\Day01-15\code\Day11\file2.py

print('读取圆周率文件判断其中是否包含自己的生日')

birth = "0311"
with open('data/pi_million_digits.txt') as f:
    lines = f.readlines()
    pi_string = ''
    for line in lines:
        pi_string += line.strip()
    if birth in pi_string:
        print('Bingo!!!')

print("------------------------------------------------------------")  # 60個

#檔案 : C:\_git\vcs\_4.python\__code\Python-100-Days-zh_TW-master\Day01-15\code\Day11\file4.py

"""
读写二进制文件
"""
import base64

with open('data/mm.jpg', 'rb') as f:
    data = f.read()
    # print(type(data))
    # print(data)
    print('字节数:', len(data))
    # 将图片处理成BASE-64编码
    print(base64.b64encode(data))

with open('tmp_girl.jpg', 'wb') as f:
    f.write(data)
print('写入完成!')

print("------------------------------------------------------------")  # 60個

#檔案 : C:\_git\vcs\_4.python\__code\Python-100-Days-zh_TW-master\Day01-15\code\Day11\json1.py

"""
读取JSON数据

"""

import json
import csv

json_str = '{"name": "骆昊", "age": 38, "title": "叫兽"}'
result = json.loads(json_str)
print(result)
print(type(result))
print(result['name'])
print(result['age'])

# 请思考如何将下面JSON格式的天气数据转换成对象并获取我们需要的信息
# 稍后我们会讲解如何通过网络API获取我们需要的JSON格式的数据
"""
    {
        "wendu": "29",
        "ganmao": "各项气象条件适宜，发生感冒机率较低。但请避免长期处于空调房间中，以防感冒。",
        "forecast": [
            {
                "fengxiang": "南风",
                "fengli": "3-4级",
                "high": "高温 32℃",
                "type": "多云",
                "low": "低温 17℃",
                "date": "16日星期二"
            },
            {
                "fengxiang": "南风",
                "fengli": "微风级",
                "high": "高温 34℃",
                "type": "晴",
                "low": "低温 19℃",
                "date": "17日星期三"
            },
            {
                "fengxiang": "南风",
                "fengli": "微风级",
                "high": "高温 35℃",
                "type": "晴",
                "low": "低温 22℃",
                "date": "18日星期四"
            },
            {
                "fengxiang": "南风",
                "fengli": "微风级",
                "high": "高温 35℃",
                "type": "多云",
                "low": "低温 22℃",
                "date": "19日星期五"
            },
            {
                "fengxiang": "南风",
                "fengli": "3-4级",
                "high": "高温 34℃",
                "type": "晴",
                "low": "低温 21℃",
                "date": "20日星期六"
            }
        ],
        "yesterday": {
            "fl": "微风",
            "fx": "南风",
            "high": "高温 28℃",
            "type": "晴",
            "low": "低温 15℃",
            "date": "15日星期一"
        },
        "aqi": "72",
        "city": "北京"
    }
"""

print("------------------------------------------------------------")  # 60個

#檔案 : C:\_git\vcs\_4.python\__code\Python-100-Days-zh_TW-master\Day01-15\code\Day11\json2.py

"""
写入JSON文件
"""

import json

teacher_dict = {'name': '白元芳', 'age': 25, 'title': '讲师'}
json_str = json.dumps(teacher_dict)
print(json_str)
print(type(json_str))
fruits_list = ['apple', 'orange', 'strawberry', 'banana', 'pitaya']
json_str = json.dumps(fruits_list)
print(json_str)
print(type(json_str))

print("------------------------------------------------------------")  # 60個


"""
创建Excel文件

"""

from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo

workbook = Workbook()
sheet = workbook.active
data = [
    [1001, '白元芳', '男', '13123456789'],
    [1002, '白洁', '女', '13233445566']
]
sheet.append(['学号', '姓名', '性别', '电话'])
for row in data:
    sheet.append(row)
tab = Table(displayName="Table1", ref="A1:E5")

tab.tableStyleInfo = TableStyleInfo(
    name="TableStyleMedium9", showFirstColumn=False,
    showLastColumn=False, showRowStripes=True, showColumnStripes=True)
sheet.add_table(tab)
workbook.save('./tmp_全班学生数据.xlsx')


print('------------------------------------------------------------')	#60個

"""
读取Excel文件

"""
from openpyxl import load_workbook
from openpyxl import Workbook

workbook = load_workbook('./data/学生明细表.xlsx')
print(workbook.sheetnames)
sheet = workbook[workbook.sheetnames[0]]
print(sheet.title)
for row in range(2, 7):
    for col in range(65, 70):
        cell_index = chr(col) + str(row)
        print(sheet[cell_index].value, end='\t')
    print()


print('------------------------------------------------------------')	#60個
"""
读取Word文件

"""

from docx import Document

doc = Document('./data/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))

print('------------------------------------------------------------')	#60個

# 處理劇本中的單一句子，去除標點並切割
def processString(str1):
    tup1 = (',', '.', ';', '?', '!', "'", '-', ':')
    strK = str1
    #Why, such is love's transgression. 要轉換成：
    #Why  such is love s transgression 然後用空白切割
    for tok in tup1:
        strK = strK.replace(tok, ' ')
    strK = strK.lower()
    slist = strK.split()
    # 串列生成，將兩個字母以上的文字以及a和i留下
    slist2 = [s for s in slist if len(s)>1 or s=='a' or s=='i']
    return slist2

filename = 'C:/_git/vcs/_4.python/_data/Romeo&Juliet.txt'

infile = open(filename, 'rt', encoding='utf-8')

freq = {}
for line in infile:
    str1 = line.strip()
    # 處理劇本中的單一句子，去除標點並切割
    slist = processString(str1)
    # 將切好的詞彙紀錄頻率
    for tok in slist:
        freq[tok] = freq.get(tok, 0)+1
infile.close()

# 開始找尋最長以及最常的詞彙
longest = ''    # 最長
maxoccur = ''    # 最常
maxv = 0
for k, v in freq.items():
    # 比出目前最長
    if len(k)>len(longest):
        longest = k
    # 比出目前最常
    if v>maxv:
        maxoccur = k
        maxv = v
        
print('最長用詞:', longest, ', 長度', len(longest))
print('最頻繁用詞:', maxoccur, ', 次數', maxv)


print('------------------------------------------------------------')	#60個

# BMP解析程式

# Little Endian的 4 bytes處理函數
def b4_2_int(b4):
    # 針對每個byte使用位元平移
    ii = b4[0] | b4[1]<<8 | b4[2]<<16 | b4[3]<<24
    return ii
# Little Endian的 2 bytes處理函數
def b2_2_int(b2):
    # 針對每個byte使用位元平移
    ii = b2[0] | b2[1]<<8
    return ii

filename = 'data/Lenna.bmp'       # 786554 bytes, 512x512
infile = open(filename, 'rb')

# 標頭識別碼：'B'和'M'
b2 = infile.read(2)
print(b2)
# 檔案大小
b4 = infile.read(4)
filesize = b4_2_int(b4)
print('File size:', filesize, 'bytes')
# 跳過不在意的欄位
infile.read(4)
# 跳過不在意的欄位
infile.read(4)
# 跳過不在意的欄位
infile.read(4)
# 圖像寬度（長度）
b4 = infile.read(4)
width = b4_2_int(b4)
# 圖像高度（寬度）
b4 = infile.read(4)
height = b4_2_int(b4)
print('Image dimension:', width, 'x', height)
b2 = infile.read(2)
planes = b2_2_int(b2)
print('Image planes:', planes)
# 圖像顏色深度（bit數）
b2 = infile.read(2)
depth = b2_2_int(b2)
print('Color depth:', depth)

infile.close()

print("------------------------------------------------------------")  # 60個

filename = 'data/DatingTestSet.txt'
stat = {}
tags = ['largeDoses', 'smallDoses', 'didntLike']
with open(filename, 'rt', encoding='utf-8') as inf:
    for line in inf:
        slist = line.strip().split('\t')
        try:
            idx = tags.index(slist[-1])
            key = tags[idx]
            stat[key] = stat.get(key, 0)+1
        except:
            # 沒出現過的tag使用index()方法
            # 會產生ValueError例外，跳過
            pass

# 使用串列生成加上sum()來找出總數
sum1 = sum([v for v in stat.values()])
for k, v in stat.items():
    # 計算百分比並完成輸出
    stat[k] = 100.0*v/sum1
    print(k+':', str(stat[k])+'%')
    
print("------------------------------------------------------------")  # 60個



print('------------------------------------------------------------')	#60個



print('------------------------------------------------------------')	#60個



print('------------------------------------------------------------')	#60個
print('作業完成')
print('------------------------------------------------------------')	#60個

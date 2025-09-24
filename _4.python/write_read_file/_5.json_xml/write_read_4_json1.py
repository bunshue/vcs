"""
#各種檔案寫讀範例 json

json常用的方法(6)

.load()   # 讀取json檔案
.loads()  # json轉字典
.dump()   # 寫進json檔案
.dumps()  # 字典轉json
load  - 將文件中的JSON數據反序列化成對象
loads - 將字符串的內容反序列化成Python對象
dump  - 將Python對象按照JSON格式序列化到文件中
dumps - 將Python對象處理成JSON格式的字符串

JSONDecoder()
JSONEncoder()

json為字串格式
json = json字串 = json_str
"""

import os
import sys
import json
import requests
import datetime
import pandas as pd

print("------------------------------------------------------------")  # 60個
print("格式轉換1, 串列轉json/json檔案")
print("------------------------------------------------------------")  # 60個

# 串列
data_list = ["鼠", "牛", "虎", "兔", "龍"]  # 串列

# 串列, 元素是字典
data_list = [
    {"英文名": "mouse", "中文名": "米老鼠", "體重": 3},
    {"英文名": "ox", "中文名": "班尼牛", "體重": 48},
    {"英文名": "tiger", "中文名": "跳跳虎", "體重": 33},
]

json_str = json.dumps(data_list)  # 串列轉json
print("串列 :", data_list)
print("串列轉json")
print("json的資料內容 :", json_str)
print("json的資料型態 :", type(json_str))

print("------------------------------------------------------------")  # 60個

# 串列
data_list = ["鼠", "牛", "虎", "兔", "龍"]  # 串列
print("data_list串列：", data_list)
print("data_list型別：", type(data_list))

json_str = json.dumps(data_list)  # 串列轉json
print("json的資料內容 :", json_str)
print("json的資料型態 :", type(json_str))


print("------------------------------------------------------------")  # 60個
print("格式轉換2, 字典轉json/json檔案")
print("------------------------------------------------------------")  # 60個

print("字典轉json")
data_dict = {"英文名": "mouse", "中文名": "米老鼠", "體重": 3}
print(type(data_dict))

json_str = json.dumps(data_dict)  # 字典轉json
print("json的資料內容 :", json_str)
print("json的資料型態 :", type(json_str))

print("json轉字典")
data_dict = json.loads(json_str)  # json轉字典
print(type(data_dict))
print(data_dict)

print("字典轉json檔案")
filename = "tmp03.json"
with open(filename, "w") as fp:
    print(type(data_dict))
    json.dump(data_dict, fp)  # 寫進json檔案, 字典轉json檔案

print("json檔案轉字典")
filename = "tmp03.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print(type(datas))
    print(datas)

json_str = json.dumps(data_dict)  # 未用排序將字典轉json  # 字典轉json
print("未用排序將字典轉換成json的物件", json_str)

json_str = json.dumps(data_dict, sort_keys=True)  # 有用排序將字典轉json  # 字典轉json
print("使用排序將字典轉換成json的物件", json_str)

json_str = json.dumps(data_dict, sort_keys=True, indent=4)  # 字典轉json
print("json的資料內容 :", json_str)

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# 字典
data_dict = {"英文名": "mouse", "中文名": "米老鼠", "體重": 3}
print("type(data_dict) :", type(data_dict))

json_str = json.dumps(data_dict, ensure_ascii=False, indent=4)  # 字典轉json
print("json的資料內容 :", json_str)

print("------------------------------------------------------------")  # 60個

print('字典轉json')

data_dict = {"英文名": "mouse", "中文名": "米老鼠", "體重": 3}
print("data_dict字典：", data_dict)
print("data_dict型別：", type(data_dict))

json_str = json.dumps(data_dict, ensure_ascii=False)  # 字典轉json
print("json的資料內容 :", json_str)
print("json的資料型態 :", type(json_str))

print("------------------------------")  # 30個

print('字典轉json')

data_dict = {"英文名": "mouse", "中文名": "米老鼠", "體重": 3}
print(json.dumps(data_dict, ensure_ascii=False))  # 字典轉json
print(json.dumps(data_dict, ensure_ascii=False, sort_keys=True))  # 字典轉json
print(json.dumps(data_dict, ensure_ascii=False, sort_keys=True, indent=4))  # 字典轉json

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# json字串 <= 字典
json_str = '{"英文名":"mouse", "中文名":"米老鼠", "體重":3}'

data_dict = json.loads(json_str)  # json轉字典
print("type(data_dict) :", type(data_dict))

print("英文名：%s" % (data_dict["英文名"]))
print("中文名：%s" % (data_dict["中文名"]))
print("體重：%d" % (data_dict["體重"]))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

print("串列轉json檔案")

# 串列
data_list = ["鼠", "牛", "虎", "兔", "龍"]  # 串列

# 串列, 元素是字典
data_list = [
    {"英文名": "mouse", "中文名": "米老鼠", "體重": 3},
    {"英文名": "ox", "中文名": "班尼牛", "體重": 48},
    {"英文名": "tiger", "中文名": "跳跳虎", "體重": 33},
]
print(type(data_list))

json_str = json.dumps(data_list)  # 串列轉json
print("json的資料內容 :", json_str)

json_str = json.dumps(
    data_list, sort_keys=True, indent=4, separators=(",", ": ")
)  # 串列轉json
print("json的資料內容 :", json_str)

data1 = json.loads(json_str)  # json轉串列
print(data1)
print(type(data1))

# 寫檔
# 串列轉json檔案
filename = "tmp01_animals.json"
# with open(filename, "w") as fp:
with open(filename, "w", encoding="utf-8") as fp:
    json.dump(data_list, fp)  # 寫進json檔案, 串列轉json檔案

filename = "tmp07a.json"
with open(filename, "w") as fp:
    json.dump(data_list, fp, ensure_ascii=False, indent=4)  # 寫進json檔案, 串列轉json檔案

filename = "tmp07b.json"
with open(filename, "w") as fp:
    json.dump(data_list, fp, indent=2, ensure_ascii=False)  # 寫進json檔案, 串列轉json檔案

filename = "tmp07c.json"
with open(filename, "w", encoding="utf-8") as fp:
    json.dump(data_list, fp, indent=2, ensure_ascii=False)  # 寫進json檔案, 串列轉json檔案

print("json檔案轉串列")
# 讀檔
filename = "tmp07c.json"
filename = "tmp01_animals.json"
# with open(filename, "r") as fp:
with open(filename, "r", encoding="utf-8") as fp:
    data_list = json.load(fp)  # 讀取json檔案
    print(type(data_list))

for data in data_list:
    for key in data:
        print(key, "：", data[key])
    print("=" * 20)

for data in data_list:
    for key in data:
        print("%s：%s" % (key, data[key]))

print(data_list)
print(type(data_list))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# json字串
json_str = '{"英文名":"mouse", "中文名":["米老鼠","傑利鼠"], "體重":3, "性別": true}'  # json字串
print("json的資料內容 :", json_str)
print("json的資料型態 :", type(json_str))

data_dict = json.loads(json_str)  # json轉字典
print("data_dict物件：", data_dict)
print("data_dict型別：", type(data_dict))
print(data_dict["英文名"])
print(data_dict["中文名"])
print(data_dict["體重"])
print(data_dict["性別"])
for key in data_dict:
    print("Key :", key, ", 內容 :", data_dict[key], ",  內容的資料型態 :", type(data_dict[key]))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個
print("格式轉換3, 元組轉json")
print("------------------------------------------------------------")  # 60個

# 元組
data_tuple = ("鼠", "牛", "虎", "兔", "龍")  # 元組
json_str = json.dumps(data_tuple)  # 元組轉json
print("元組轉json :", json_str)
print("json的資料型態 :", type(json_str))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# json字串 <= 串列, 元素是字典
json_str = """
    [
    {"英文名": "mouse", "中文名": "米老鼠", "體重": 3},
    {"英文名": "ox", "中文名": "班尼牛", "體重": 48},
    {"英文名": "tiger", "中文名": "跳跳虎", "體重": 33}
    ]
"""
print("json的資料型態 :", type(json_str))

data_list = json.loads(json_str)  # json轉串列
print(type(data_list))

for data in data_list:
    print("英文名：%s" % (data["英文名"]))
    print("中文名：%s" % (data["中文名"]))
    print("體重：%d" % (data["體重"]))
    print("=" * 20)

for data in data_list:
    for key in data:
        print("%s：%s" % (key, data[key]))
    print("=" * 20)


print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

"""
filename = "tmp18.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
print(type(datas))
print(datas)

key = "aaa"
datas[key] = "AAAA"
with open(filename, "w") as fp:
    print(type(datas))
    json.dump(datas, fp)  # 寫進json檔案

key = "bbb"
datas[key] = "BBBB"
with open(filename, "w") as fp:
    print(type(datas))
    json.dump(datas, fp)  # 寫進json檔案

key = "ccc"
datas[key] = "CCCC"
with open(filename, "w") as fp:
    print(type(datas))
    json.dump(datas, fp)  # 寫進json檔案

print("存了3筆資料")

print("讀取資料")
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
print(datas)

key = "bbb"
print(datas[key])
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

url = "https://fchart.github.io/json/GoogleBooks.json"
filename = "tmp20_Books.json"
r = requests.get(url)
r.encoding = "utf8"
data_dict = json.loads(r.text)  # json轉字典
with open(filename, "w") as fp:
    print(type(data_dict))
    json.dump(data_dict, fp)  # 寫進json檔案, 字典轉json檔案

print("------------------------------------------------------------")  # 60個
print("解讀json專區 ST")
print("------------------------------------------------------------")  # 60個

filename = "data/Student.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print(type(datas))
    print(datas)

json_str = json.dumps(datas)  # 字典轉json
print("json的資料內容 :", json_str)

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

filename = "data/populations.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")
    # print(datas)
    i = 0
    for data in datas:
        # print(data)  # 字典格式
        y = data["Year"]
        n = data["Country Name"]  # 國家名稱
        c = data["Country Code"]  # 國家代碼
        p = int(float(data["Numbers"]))  # 人口數
        print(y, "/", n, "/", c, "/", p)
        i += 1
        if i > 10:
            break

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

filename = "data/aqi.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")
    # print(datas)
    i = 0
    for data in datas:
        # print(data)  # 字典格式
        county = data["County"]  # 城市名稱
        sitename = data["SiteName"]  # 站台名稱
        siteid = data["SiteId"]  # 站台ID
        pm25 = data["PM2.5"]  # PM2.5值
        print(
            "城市名稱 =%4s  站台ID =%3s  PM2.5值 =%3s  站台名稱 = %s "
            % (county, siteid, pm25, sitename)
        )
        i += 1
        if i > 10:
            break

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

filename = "data/eq_data_1_day_m1.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")
    # print(datas)
    for data in datas:
        print(data)

print("type :")
print(datas["type"])

print("features :")
print(datas["features"])
feature_data = datas["features"]
i = 0
for feature in feature_data:
    mag = feature["properties"]["mag"]
    lon = feature["geometry"]["coordinates"][0]
    lat = feature["geometry"]["coordinates"][1]
    print(mag, "/", lon, "/", lat)
    i += 1
    if i > 10:
        break

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

filename = "D:/_git/vcs/_1.data/______test_files1/_json/data_earthquake.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")

print("過去7天全球發生重大的地震資訊：")
for data in datas["features"]:
    print("地點:{}".format(data["properties"]["place"]))
    print("震度:{}".format(data["properties"]["mag"]))
    et = float(data["properties"]["time"]) / 1000.0
    d = datetime.datetime.fromtimestamp(et).strftime("%Y-%m-%d %H:%M:%S")
    print("時間:{}".format(d))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

filename = "D:/_git/vcs/_1.data/______test_files1/_json/ChinaBoundary_Province_City"

with open(filename, "r", encoding="utf-8") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")
    print(datas["ID"])
    print(datas["code"])
    print(datas["name"])
    # print(datas['Province'])
    print(datas["Province"][0]["ID"])
    print(datas["Province"][0]["code"])
    print(datas["Province"][0]["name"])
    # print(datas['Province'][0]['rings'])
    # print(datas['Province'][0]['City'])
    print(datas["Province"][0]["City"][0]["code"])
    print(datas["Province"][0]["City"][0]["name"])

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

print("讀取一個JSON字典")

from difflib import get_close_matches

filename = "data/WORD_DICTIONARY.json"
with open(filename, "r") as fp:
    datas = json.load(fp)  # 讀取json檔案
    print("共有 :", len(datas), "筆資料")
    # print(datas)

text = "XYZ"  # 找不到
text = "MoisTxure"  # 不完全相符
text = "MOISTURE"  # 全相符
text = "abiotic factor"  # 全相符

w = text.lower()
if w in datas:
    print("全相符")
    for i in range(len(datas[w])):
        axe = str(datas[w][i])
        print(axe)
        # text1.insert(END,w+" : "+axe+"\n")
        # speak(axe)
elif len(get_close_matches(w, datas.keys())) > 0:
    print("不完全相符, 最接近的字 %s " % get_close_matches(w, datas.keys())[0])
    for i in range(len(datas[get_close_matches(w, datas.keys())[0]])):
        print(
            str(get_close_matches(w, datas.keys())[0])
            + " : "
            + str(datas[get_close_matches(w, datas.keys())[0]][i])
            + "\n"
        )
else:
    print("找不到")


print("------------------------------------------------------------")  # 60個
print("解讀json專區 SP")
print("------------------------------------------------------------")  # 60個

import matplotlib.pyplot as plt

# 讀取109年9月臺中市10大易肇事路口.json資料並放入listTrafficEvent串列物件

filename = "data/109年9月臺中市10大易肇事路口.json"
with open(filename, "r", encoding="utf-8") as fp:
    listTrafficEvent = json.load(fp)  # 讀取json檔案

# 將 listTrafficEvent 串列中的每筆字典物件印出來
for trafficEvent in listTrafficEvent:
    for key in trafficEvent:
        show = False
        # 若該筆字典的編號(key)有值(value)即印出
        if trafficEvent["編號"] != "":
            print("%s：%s" % (key, trafficEvent[key]))
            show = True
        else:
            break
    if show == True:
        print("=" * 30)

total = 0.0  # 統計肇事數量
listAllEvent = []  # listAllEvent用來存放車禍主要肇因
for trafficEvent in listTrafficEvent:
    if trafficEvent["編號"] != "":
        listAllEvent += [trafficEvent["主要肇因"]]
        total += 1

listEvent = []  # 存放所有車禍主要肇因，此串列存放不重複的車禍主要肇因
listCount = []  # 存放各主要肇因對應的車禍數量
for event in set(listAllEvent):  # 使用set()移除listAllEvent串列中重複的主要肇因
    print("主要肇因 %s 共 %s 件" % (event, listAllEvent.count(event)))
    listEvent += [event]
    listCount += [listAllEvent.count(event)]

# 計算各車禍主要肇因的百分比，並放入listPercent串列
listPercent = []
for c in listCount:
    listPercent.append(round((float(c) / total) * 100))

# 繪製車禍主要肇因的圓餅圖
font = {"family": "DFKai-SB"}
plt.rc("font", **font)
plt.pie(listPercent, labels=listEvent, autopct="%3.1f%%")
# plt.show()  # 顯示圓餅圖

print("圖表建置成成功")

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# 抓取地震資料 用json拆解

url = "http://earthquake.usgs.gov/earthquakes/feed/v1.0/summary/4.5_week.geojson"
html = requests.get(url)

datas = json.loads(html.text)  # json轉字典
print("共有 :", len(datas), "筆資料")

dataset = list()

for data in datas["features"]:
    item = dict()
    eptime = float(data["properties"]["time"]) / 1000.0
    d = datetime.datetime.fromtimestamp(eptime).strftime("%Y-%m-%d %H:%M:%S")
    item["eqtime"] = d
    item["mag"] = data["properties"]["mag"]
    item["place"] = data["properties"]["place"]
    dataset.append(item)
    # print(item)

print("------------------------------")  # 30個

for data in dataset:
    data = '(`eqtime`,`mag`,`place`) values("{}",{},"{}");'.format(
        data["eqtime"], data["mag"], data["place"]
    )
    # print(data)


print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

"""
#import ast

data = dict()
with open('data\password.txt','r', encoding = 'UTF-8-sig') as fp:
   filedata = fp.read()
   print(filedata)
#   filedata = filedata.replace("\'", "\"")
#   data = ast.literal_eval(filedata)
   data = json.loads(filedata)

print(type(data),data)
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個


from collections import defaultdict


def print_scores(filename):
    with open(filename, "r") as fp:
        datas = json.load(fp)  # 讀取json檔案
        result = defaultdict(list)

        print("班級:", datas["class"])
        for data in datas["score"]:
            for subject, score in data.items():
                result[subject].append(score)

        for subject, scores in result.items():
            print("科目:", subject)
            print("\t最高分:", max(scores))
            print("\t最低分:", min(scores))
            print("\t平均:", sum(scores) / len(scores))


# 讀取 JSON 檔
print_scores(r".\data\score.json")


def print_dir_scores(dirname):
    for filename in os.listdir(dirname):
        if filename.endswith(".json"):
            print("讀取檔案: ", filename)
            print_scores(os.path.join(dirname, filename))


# 批次檔案讀取
print_dir_scores(r".\data\scores")

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

from collections import OrderedDict

print("字典轉json")

# 字典
data_dict = {"英文名": "mouse", "中文名": "米老鼠", "體重": 3}
print(type(data_dict))

json_str = json.dumps(data_dict)  # str类型

data_dict = json.loads(json_str)

# 寫json檔案
with open("tmp_data.json", "w") as f:
    json.dump(data_dict, f)

# 讀json檔案
with open("tmp_data.json", "r") as f:
    data_dict = json.load(f)

# 使用object_pairs_hook
# json字串 <= 字典
json_str = '{"英文名":"mouse", "中文名":"米老鼠", "體重":3}'
print(type(json_str))

data_dict = json.loads(json_str, object_pairs_hook=OrderedDict)
print(data_dict)

# 解码为自定义对象
data_dict = json.loads(json_str)

print(json.dumps(data_dict))
print(json.dumps(data_dict, indent=4))

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個
print("作業完成")
print("------------------------------------------------------------")  # 60個
sys.exit()


print("------------------------------------------------------------")  # 60個

# 3030
print("------------------------------")  # 30個

import pprint as pp

# data = xxxx json格式

pp.pprint(data)
# print(data)

print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

"""
response = requests.post(xxxx)
result = response.json()
print(result[0]['translations'][0]['text'])
print(result)	#列印結果
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

# json轉dict

from urllib.request import urlopen
from html.entities import html5

entities_url = "http://dev.w3.org/html5/spec/entities.json"


def get_json(url):
    # Download the json file from the url and returns a decoded object.
    with urlopen(url) as fp:
        data = fp.read().decode("utf-8")
    return json.loads(data)


def create_dict(entities):
    # Create the html5 dict from the decoded json object.
    new_html5 = {}
    for name, value in entities.items():
        new_html5[name.lstrip("&")] = value["characters"]
    return new_html5


new_html5 = create_dict(get_json(entities_url))


print("------------------------------------------------------------")  # 60個
"""
postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 查看有哪些欄位
print(data.keys())
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 查看有哪些內容
print(df)
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

"""
postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df.loc["title", 0])
print()
print(">>>>>原Po文章")
print(str(df.loc["content", 0]).strip())
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個
"""
from datetime import datetime

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df[0]["title"])
print()
print(">>>>>原Po文章\n")
print(str(df[0]["content"]).strip())
url = "https://www.dcard.tw/service/api/v2/posts/" + postid + "/comments?popular=True"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 取出前3筆資料
for i in range(3):
    print("*******熱門留言" + str(i) + ":")
    print(df.loc[i, "content"])
    # 將最後修改文字日期轉換成日期
    updatedate = datetime.fromisoformat(str(df.loc[i, "updatedAt"]).replace("Z", ""))
    # 計算現在和最後修改日期的時間差
    datediff = datetime.today() - updatedate
    # 顯示幾天前有最新留言
    print(datediff.days, "天前")
    print()
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個
"""
from datetime import datetime

postid = "237123381"
url = "https://www.dcard.tw/service/api/v2/posts/" + postid
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為只有一欄，而且沒有欄位名稱
# 否則只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame.from_dict(data, orient="index")
# 取出title、content內容
print(">>>>>文章標題")
print(df[0]["title"])
print()
print(">>>>>原Po文章\n")
print(str(df[0]["content"]).strip())
url = "https://www.dcard.tw/service/api/v2/posts/" + postid + "/comments"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 除了updatedAt, content, likeCount三個欄位以外，全部刪除
# 刪除欄依定要指定axis=1
# inplace=True真實刪除
df.drop(
    [
        "id",
        "anonymous",
        "postId",
        "createdAt",
        "floor",
        "withNickname",
        "hiddenByAuthor",
        "meta",
        "gender",
        "school",
        "host",
        "reportReason",
        "mediaMeta",
        "hidden",
        "inReview",
        "reportReasonText",
        "isSuspiciousAccount",
        "isModerator",
        "doorplate",
        "edited",
        "postAvatar",
        "activityAvatar",
        "verifiedBadge",
        "memberType",
        "enablePrivateMessage",
        "department",
    ],
    inplace=True,
    axis=1,
)
# 依據likeCout內容排序，以降序排列
df.sort_values(by="likeCount", inplace=True, ascending=False)
# 要記得重設index，這樣df.loc[]的結果才會正確
df = df.reset_index(drop=True)
# 取出前3筆資料
for i in range(3):
    print("*******熱門留言" + str(i) + ":")
    print(df.loc[i, "content"])
    # 將最後修改文字日期轉換成日期
    updatedate = datetime.fromisoformat(str(df.loc[i, "updatedAt"]).replace("Z", ""))
    # 計算現在和最後修改日期的時間差
    datediff = datetime.today() - updatedate
    # 顯示幾天前有最新留言
    print(datediff.days, "天前")
    print()
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)
"""
print("------------------------------------------------------------")  # 60個
print("------------------------------------------------------------")  # 60個

url = "https://od.cdc.gov.tw/eic/Day_Confirmation_Age_County_Gender_19CoV.json"
res = requests.get(url)
# 有兩種方法，下面兩行任選一種都可以
# data = json.loads(res.text) #方法1
data = res.json()  # 方法2
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 轉換'確定病例數'欄位內容為整數(以利後面加總)
df["確定病例數"] = df["確定病例數"].astype(int)

startdate = "2021/07/20"
# startdate=input("請輸日期(yyyy/mm/dd)")
# 只保留個案研判日、是否為境外移入、確定病例數 三個欄位
df.drop(["確定病名", "縣市", "鄉鎮", "性別", "年齡層"], inplace=True, axis=1)
# 設定過濾條件(累積)
indexNames = df[df["個案研判日"] > startdate].index
# 刪除所有大於指定日期的個案
df.drop(indexNames, inplace=True)
# 計算累積確診人數
total = df.sum()["確定病例數"]
print("累積確診人數:", total)
# 設定過濾條件(今日)
indexNames = df[df["個案研判日"] != startdate].index
# 刪除所有不是指定日期的個案
df.drop(indexNames, inplace=True)
# 計算境外移入人數
imported = df[df["是否為境外移入"] == "是"].sum()["確定病例數"]
print("今日境外移入人數:", imported)
# 計算本土人數
domestic = df[df["是否為境外移入"] == "否"].sum()["確定病例數"]
print("今日本土人數:", domestic)
total = imported + domestic
print("今日總人數:", total)
# 若要查看有哪些欄位
print((df.keys()))
# 也可以透過df.info()查看
df.info()
# 也可以直接看前幾列
df.head(3)

print("------------------------------------------------------------")  # 60個

import csv
from docx import Document

url = "https://od.cdc.gov.tw/eic/Day_Confirmation_Age_County_Gender_19CoV.json"
res = requests.get(url)
data = res.json()
# 因為有欄位名稱
# 只要 df = pd.DataFrame(data) 即可
df = pd.DataFrame(data)
# 轉換'確定病例數'欄位內容為整數(以利後面加總)
df["確定病例數"] = df["確定病例數"].astype(int)
startdate = "2021/07/20"
startdate = input("請輸日期(yyyy/mm/dd)")
# 只保留個案研判日、是否為境外移入、確定病例數 三個欄位
df.drop(["確定病名", "縣市", "鄉鎮", "性別", "年齡層"], inplace=True, axis=1)
# 設定過濾條件(累積)
indexNames = df[df["個案研判日"] > startdate].index
# 刪除所有大於指定日期的個案
df.drop(indexNames, inplace=True)
# 計算累積確診人數
total = df.sum()["確定病例數"]
print("累積確診人數:", total)
# 設定過濾條件(今日)
indexNames = df[df["個案研判日"] != startdate].index
# 刪除所有不是指定日期的個案
df.drop(indexNames, inplace=True)
# 計算境外移入人數
imported = df[df["是否為境外移入"] == "是"].sum()["確定病例數"]
print("今日境外移入人數:", imported)
# 計算本土人數
domestic = df[df["是否為境外移入"] == "否"].sum()["確定病例數"]
print("今日本土人數:", domestic)
print("今日總人數:", domestic + imported)

url = "https://od.cdc.gov.tw/eic/covid19/covid19_tw_stats.csv"
df = pd.read_csv(url)
# 只保留確診、死亡、送驗、排除四個欄位
df.drop(["昨日確診", "昨日排除", "昨日送驗"], inplace=True, axis=1)
totaldeath = df.loc[0, "死亡"]
print("累積死亡人數", totaldeath)

# 開啟指定的Word檔案
my_doc = Document("指揮中心快訊.docx")
# 設定要取代的串列list
replacements = {
    "%Date%": startdate,
    "%N1%": str(domestic + imported),
    "%N2%": str(domestic),
    "%N3%": str(imported),
    "%N4%": str(totaldeath),
    "%N5%": str(total),
}
# 尋找要取代的關鍵字
for key in replacements:
    # 尋找所有的表格
    for table in my_doc.tables:
        # 尋找表格中的列
        for row in table.rows:
            # 尋找列中的每個儲存格
            for cell in row.cells:
                # 尋找儲存格中的每一個斷落
                for paragraph in cell.paragraphs:
                    # 如果關鍵字出現在段落中
                    if key in paragraph.text:
                        # 為了避免修改掉原有的樣式，必須用這個方法處理
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                text = inline[i].text.replace(key, replacements[key])
                                inline[i].text = text
# 指定儲存的檔案名稱
fname = "指揮中心快訊" + startdate.replace("/", "") + ".docx"
# 儲存檔案
my_doc.save(fname)

"""
請輸日期(yyyy/mm/dd)2021/07/20
累積確診人數: 15450
今日境外移入人數: 6
今日本土人數: 18
今日總人數: 24
累積死亡人數 846
"""

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個

"""
# 新增員工記錄函式
def fnCreate():
    uid = input("編號：")
    if uid in listUid:
        print("編號重複，無法在記憶體中新增員工記錄")
        return
    name = input("姓名︰")
    salary = int(input("薪資︰"))
    newMember = {"編號": uid, "姓名": name, "薪資": salary}
    newMember = {"編號": uid, "姓名": name, "薪資": salary}

            listMember.remove(member)
            listUid.remove(uid)

filename = "MemberInfo.json"

    with open(filename, "w", encoding="utf-8") as fp:
        print(type(listMember))
        json.dump(listMember, fp, ensure_ascii=False, indent=4)  # 寫進json檔案
        print("記憶體中的員工記錄成功儲存至 %s 檔案" % (filename))

if os.path.exists(filename):
    filename = "MemberInfo.json"
    with open(filename, "r", encoding="utf-8") as fp:
        listMember = json.load(fp)  # 讀取json檔案
"""


"""
字典
data_dict = json.loads(data)  # json 轉 字典
字串
data_str

串列
data_list

json為字串格式
json = json字串 = json_str

json字串 = 串列 用''包起來變成json字串
json字串 = 字典 用''包起來變成json字串
寫檔讀檔後，會依原本儲存的格式得到格式

"""


"""
#各種檔案寫讀範例 docx 1


pip install python-docx

"""

print('------------------------------------------------------------')	#60個

import docx

filename = 'data/python_docx1.docx'

doc = docx.Document(filename)
for p in doc.paragraphs:
    print(p.text)

print('------------------------------------------------------------')	#60個

#讀取excel檔, 製作word檔

import openpyxl
from docxtpl import DocxTemplate
import datetime

filename = 'data/python_ReadWrite_EXCEL6_student_data2.xlsx'

workbook = openpyxl.load_workbook(filename)
sheet = workbook.active

list_values = list(sheet.values)
print(list_values)

filename = 'data/python_ReadWrite_WORD1_certificate.docx'

# Generate docs
doc = DocxTemplate(filename)

for value_tuple in list_values[1:]:
    doc.render({"name": value_tuple[0],
                "course": value_tuple[1],
                "date": value_tuple[2],
                "instructor":value_tuple[3]})
    
    doc_name = "tmp_certificate" + value_tuple[0] + value_tuple[1] + ".docx"
    print(doc_name)
    doc.save(doc_name)



print('------------------------------------------------------------')	#60個

"""
讀取word檔, 製作word檔
invoice : 發票，發貨單
"""

import datetime
from docxtpl import DocxTemplate

filename1 = 'data/python_ReadWrite_WORD2_invoice_template.docx'

doc = DocxTemplate(filename1)

invoice_list = [[2, "pen", 0.5, 1],
                [1, "paper pack", 5, 5],
                [2, "notebook", 2, 4]]

"""
doc.render({"name":"john", 
            "phone":"555-55555",
            "invoice_list": invoice_list,
            "subtotal":10,
            "salestax":"10%",
            "total":9})
"""
name = 'David Wang'
phone = '0912-345678'
subtotal = sum(item[3] for item in invoice_list) 
salestax = 0.1
total = subtotal*(1-salestax)

doc.render({"name":name,
            "phone":phone,
            "invoice_list": invoice_list,
            "subtotal":subtotal,
            "salestax":str(salestax*100)+"%",
            "total":total})

filename2 = "tmp_new_invoice" + name + datetime.datetime.now().strftime("%Y-%m-%d-%H%M%S") + ".docx"
doc.save(filename2)


print('------------------------------------------------------------')	#60個

#讀取Word文件

from docx import Document

doc = Document('./data/用函數還是用復雜的表達式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))


print('------------------------------------------------------------')	#60個



print('------------------------------------------------------------')	#60個
print('作業完成')
print('------------------------------------------------------------')	#60個





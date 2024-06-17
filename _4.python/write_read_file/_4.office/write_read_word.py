"""
讀寫 Word 檔案, 使用 docx
pip install python-docx

"""

import sys

import docx

print("------------------------------------------------------------")  # 60個

filename_r = "data/python_docx1.docx"

doc = docx.Document(filename_r)
for p in doc.paragraphs:
    print(p.text)

print("------------------------------------------------------------")  # 60個

# 讀取excel檔, 製作word檔

import openpyxl
from docxtpl import DocxTemplate
import datetime

filename_r1 = "data/python_ReadWrite_EXCEL6_student_data2.xlsx"
workbook = openpyxl.load_workbook(filename_r1)
sheet = workbook.active
list_values = list(sheet.values)
print(list_values)

filename_r2 = "data/python_ReadWrite_WORD1_certificate.docx"
# Generate docs
doc = DocxTemplate(filename_r2)
for value_tuple in list_values[1:]:
    doc.render(
        {
            "name": value_tuple[0],
            "course": value_tuple[1],
            "date": value_tuple[2],
            "instructor": value_tuple[3],
        }
    )

    doc_name = "tmp_certificate" + value_tuple[0] + value_tuple[1] + ".docx"
    print(doc_name)
    doc.save(doc_name)

print("------------------------------------------------------------")  # 60個

"""
讀取word檔, 製作word檔
invoice : 發票，發貨單
"""

import datetime
from docxtpl import DocxTemplate

filename_r = "data/python_ReadWrite_WORD2_invoice_template.docx"

doc = DocxTemplate(filename_r)

invoice_list = [[2, "pen", 0.5, 1], [1, "paper pack", 5, 5], [2, "notebook", 2, 4]]

"""
doc.render({"name":"john", 
            "phone":"555-55555",
            "invoice_list": invoice_list,
            "subtotal":10,
            "salestax":"10%",
            "total":9})
"""
name = "David Wang"
phone = "0912-345678"
subtotal = sum(item[3] for item in invoice_list)
salestax = 0.1
total = subtotal * (1 - salestax)

doc.render(
    {
        "name": name,
        "phone": phone,
        "invoice_list": invoice_list,
        "subtotal": subtotal,
        "salestax": str(salestax * 100) + "%",
        "total": total,
    }
)

filename_w = (
    "tmp_new_invoice"
    + name
    + datetime.datetime.now().strftime("%Y-%m-%d-%H%M%S")
    + ".docx"
)
doc.save(filename_w)

print("------------------------------------------------------------")  # 60個

# 讀取Word文件

from docx import Document

filename_r = "./data/用函數還是用復雜的表達式.docx"

doc = Document(filename_r)
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print("".join(content))

print("------------------------------------------------------------")  # 60個

from docx import Document #引入套件
from docx.shared import Cm
from docx.shared import Inches

#開新的Word檔案
document = Document()  
#標題 段落樣式
document.add_heading('標題 段落樣式',0) 
#內文 段落樣式
p = document.add_paragraph('內文 段落樣式包含一些 ') 
p.add_run('粗體').bold = True
p.add_run(' 還有一些 ')
p.add_run('斜體字.').italic = True
#標題1 段落樣式
document.add_heading('標題1 段落樣式', level=1) 
#鮮明引文 段落樣式
document.add_paragraph('鮮明引文 段落樣式', style='Intense Quote') 
#項目符號 段落樣式
document.add_paragraph('項目符號 段落樣式,1', style='List Bullet')
document.add_paragraph('項目符號 段落樣式,2', style='List Bullet')
document.add_paragraph('清單號碼 段落樣式,1', style='List Number')
document.add_paragraph('清單號碼 段落樣式,2', style='List Number')
#插入圖片，指定寬度5cm
document.add_picture('data/Google-Colab-Guide-1024x683.jpg' , width=Cm(5))
#插入表格，先增加1列標題列
table = document.add_table(rows=1, cols=3, style = 'Table Grid') 
#表格標題列
header_cells = table.rows[0].cells 
header_cells[0].text = '班級'
header_cells[1].text = '座號'
header_cells[2].text = '姓名'
#表格內容第1列
row_cells = table.add_row().cells 
row_cells[0].text = '101'
row_cells[1].text = '01'
row_cells[2].text = '劉的華'
#表格內容第2列
row_cells = table.add_row().cells 
row_cells[0].text = '101'
row_cells[1].text = '02'
row_cells[2].text = '郭付錢'
#換頁符號
document.add_page_break()
#儲存檔案 
document.save('tmp_word1.docx') 

print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個


print("------------------------------------------------------------")  # 60個
print("作業完成")
print("------------------------------------------------------------")  # 60個
